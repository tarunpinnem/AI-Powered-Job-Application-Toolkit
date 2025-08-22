#!/usr/bin/env python3
"""
AI Career Success Platform
Enterprise-grade career development toolkit with AI-powered features
"""

# Import required libraries
import os
import re
import json
import time
import logging
import sqlite3
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
from collections import defaultdict
import hashlib

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Flask imports
from flask import Flask, render_template, request, jsonify, session, send_file
from flask_cors import CORS

# Optional imports with fallbacks
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    HAS_LIMITER = True
except ImportError:
    HAS_LIMITER = False

# AI and NLP imports (with fallbacks for demo)
try:
    import openai
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False
    
try:
    import spacy
    HAS_SPACY = True
except ImportError:
    HAS_SPACY = False

try:
    import nltk
    HAS_NLTK = True
except ImportError:
    HAS_NLTK = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    import redis
    HAS_REDIS = True
except ImportError:
    HAS_REDIS = False

# Environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()  # Load environment variables from .env file
    HAS_DOTENV = True
except ImportError:
    HAS_DOTENV = False

# Document parsing imports
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class CareerSuccessPlatform:
    """
    Main application class for the AI Career Success Platform
    """
    
    def __init__(self):
        self.app = Flask(__name__)
        self.app.secret_key = os.getenv('SECRET_KEY', 'dev-key-change-in-production')
        
        # Initialize components
        self.setup_cors()
        self.setup_rate_limiting()
        self.setup_database()
        self.setup_ai_services()
        self.setup_routes()
        
        logger.info("🚀 AI Career Success Platform initialized successfully")
    
    def setup_cors(self):
        """Configure CORS for API access"""
        CORS(self.app, resources={
            r"/api/*": {
                "origins": ["http://localhost:3000", "http://127.0.0.1:5000"],
                "methods": ["GET", "POST", "PUT", "DELETE"],
                "allow_headers": ["Content-Type", "Authorization"]
            }
        })
    
    def setup_rate_limiting(self):
        """Setup rate limiting if available"""
        if HAS_LIMITER:
            self.limiter = Limiter(
                app=self.app,
                key_func=get_remote_address,
                default_limits=["200 per day", "50 per hour"]
            )
            logger.info("✅ Rate limiting enabled")
        else:
            self.limiter = None
            logger.warning("⚠️ Rate limiting disabled (flask-limiter not available)")
    
    def setup_database(self):
        """Initialize SQLite database"""
        self.db_path = 'career_platform.db'
        self.init_database()
        logger.info("✅ Database initialized")
    
    def setup_ai_services(self):
        """Initialize AI services"""
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        
        if HAS_OPENAI and self.openai_api_key:
            # Initialize OpenAI client (new API format)
            try:
                from openai import OpenAI
                self.openai_client = OpenAI(api_key=self.openai_api_key)
                logger.info("✅ OpenAI GPT-4 integration enabled")
            except Exception as e:
                logger.warning(f"⚠️ OpenAI client initialization failed: {e}")
                self.openai_client = None
        else:
            logger.warning("⚠️ OpenAI integration disabled (API key or library not available)")
            self.openai_client = None
        
        # Initialize NLTK data if available
        if HAS_NLTK:
            try:
                import nltk
                nltk.download('punkt', quiet=True)
                nltk.download('stopwords', quiet=True)
                nltk.download('wordnet', quiet=True)
                logger.info("✅ NLTK initialized")
            except Exception as e:
                logger.warning(f"⚠️ NLTK initialization failed: {e}")
    
    def init_database(self):
        """Create database tables"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute('''
                CREATE TABLE IF NOT EXISTS analyses (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    session_id TEXT,
                    resume_text TEXT,
                    job_description TEXT,
                    industry TEXT,
                    overall_score INTEGER,
                    analysis_data TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            conn.execute('''
                CREATE TABLE IF NOT EXISTS user_sessions (
                    session_id TEXT PRIMARY KEY,
                    user_data TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    last_active TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            conn.execute('''
                CREATE TABLE IF NOT EXISTS analytics (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    event_type TEXT,
                    event_data TEXT,
                    session_id TEXT,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''')
            
            conn.commit()
    
    def setup_routes(self):
        """Setup all application routes"""
        
        @self.app.route('/')
        def index():
            """Main application page"""
            return render_template('career_platform.html')
        
        @self.app.route('/test-upload')
        def test_upload_page():
            """Test upload page for debugging"""
            return render_template('test_upload.html')
        
        @self.app.route('/upload-help')
        def upload_help_page():
            """Upload help and instructions page"""
            return render_template('upload_help.html')

        @self.app.route('/api/upload-help')
        def upload_help():
            """Provide helpful instructions for file uploads"""
            return jsonify({
                'success': True,
                'instructions': {
                    'file_upload': {
                        'supported_formats': ['PDF', 'DOCX', 'TXT'],
                        'max_size': '10MB',
                        'tips': [
                            'Ensure your PDF is not password protected',
                            'Use standard fonts in your documents',
                            'Avoid heavily formatted or image-based resumes',
                            'Make sure your document contains actual text (not just images)'
                        ]
                    },
                    'text_paste': {
                        'how_to': [
                            'Open your resume in any text editor or word processor',
                            'Select all text (Ctrl+A or Cmd+A)',
                            'Copy the text (Ctrl+C or Cmd+C)',
                            'Paste into the text area on our platform'
                        ],
                        'what_not_to_paste': [
                            'Do not paste raw PDF content or binary data',
                            'Avoid pasting formatted content that includes special characters',
                            'Do not copy from PDF viewers if the text appears garbled'
                        ]
                    },
                    'troubleshooting': [
                        'If file upload fails, try converting to plain text format',
                        'If PDF extraction fails, try opening the PDF in a different viewer and copying the text',
                        'For best results, use a clean, well-formatted Word document or plain text file',
                        'Contact support if you continue experiencing issues'
                    ]
                }
            })

        @self.app.route('/health')
        def health_check():
            """Health check endpoint"""
            status = {
                'status': 'healthy',
                'timestamp': datetime.now().isoformat(),
                'version': '1.0.0',
                'features': {
                    'openai': bool(self.openai_client),
                    'nltk': HAS_NLTK,
                    'spacy': HAS_SPACY,
                    'rate_limiting': HAS_LIMITER,
                    'redis': HAS_REDIS
                }
            }
            return jsonify(status)
        
        @self.app.route('/api/test-upload', methods=['POST'])
        def test_upload():
            """Test file upload functionality"""
            try:
                logger.info("Test upload endpoint called")
                
                if 'resume_file' not in request.files:
                    logger.warning("No file in request.files")
                    return jsonify({
                        'success': False,
                        'error': 'No file uploaded',
                        'debug_info': {
                            'files_in_request': list(request.files.keys()),
                            'form_data': dict(request.form)
                        }
                    }), 400
                
                file = request.files['resume_file']
                if not file or not file.filename:
                    logger.warning("Empty file or no filename")
                    return jsonify({
                        'success': False,
                        'error': 'No file selected'
                    }), 400
                
                filename = file.filename
                logger.info(f"Processing file: {filename}")
                
                # Get file size
                file.seek(0, 2)  # Seek to end
                file_size = file.tell()
                file.seek(0)  # Reset to beginning
                
                logger.info(f"File size: {file_size} bytes")
                
                # Test extraction with detailed logging
                logger.info("Starting text extraction...")
                extracted_text = self.extract_text_from_uploaded_file(file)
                logger.info(f"Extraction completed. Text length: {len(extracted_text) if extracted_text else 0}")
                
                return jsonify({
                    'success': True,
                    'filename': filename,
                    'file_size': file_size,
                    'extracted_length': len(extracted_text) if extracted_text else 0,
                    'extraction_preview': extracted_text[:300] if extracted_text else "No text extracted",
                    'extracted_text': extracted_text,  # Full text for main page usage
                    'debug_info': {
                        'has_pdfplumber': HAS_PDFPLUMBER,
                        'has_pypdf2': HAS_PDF,
                        'has_docx': HAS_DOCX,
                        'file_extension': os.path.splitext(filename)[1].lower()
                    }
                })
                
            except Exception as e:
                logger.error(f"Test upload error: {e}", exc_info=True)
                return jsonify({
                    'success': False,
                    'error': str(e),
                    'debug_info': {
                        'error_type': type(e).__name__
                    }
                }), 500

        @self.app.route('/api/analyze-resume', methods=['POST'])
        def analyze_resume():
            """Comprehensive resume analysis endpoint"""
            try:
                # Handle file upload if present
                resume_text = ""
                file_upload_attempted = False
                
                if 'resume_file' in request.files:
                    file = request.files['resume_file']
                    if file and file.filename:
                        file_upload_attempted = True
                        logger.info(f"File upload detected: {file.filename}")
                        resume_text = self.extract_text_from_uploaded_file(file)
                        logger.info(f"Extracted text length: {len(resume_text) if resume_text else 0}")
                
                # Also get text input (in case user typed or pasted)
                if not resume_text:
                    form_text = request.form.get('resume_text', '').strip()
                    if form_text:
                        logger.info(f"Using form text, length: {len(form_text)}")
                        resume_text = form_text
                    else:
                        logger.warning("No resume text found in form or file upload")
                
                job_description = request.form.get('job_description', '').strip()
                industry = request.form.get('industry', 'technology')
                
                if not resume_text:
                    if file_upload_attempted:
                        return jsonify({
                            'success': False,
                            'error': 'Unable to extract text from the uploaded file. Please try: 1) Copy and paste your resume text directly into the text area, 2) Convert your PDF to a Word document (.docx), 3) Save as plain text (.txt), or 4) Check our upload help guide.',
                            'help_url': '/api/upload-help'
                        }), 400
                    else:
                        return jsonify({
                            'success': False,
                            'error': 'Resume text is required. Please upload a file or paste your resume content in the text area.'
                        }), 400
                
                # Check if resume text is actually readable content
                if len(resume_text.strip()) < 50:
                    return jsonify({
                        'success': False,
                        'error': 'Resume text is too short. Please provide a complete resume with at least 50 characters.'
                    }), 400
                
                # Check if the text looks like resume content
                # Resume should contain some common words/patterns
                resume_indicators = 0
                common_resume_words = ['experience', 'education', 'skills', 'work', 'job', 'degree', 'university', 'company', 'project', 'responsibility', 'employment', 'career', 'professional', 'manager', 'engineer', 'developer', 'analyst', 'coordinator', 'specialist', 'director']
                for word in common_resume_words:
                    if word.lower() in resume_text.lower():
                        resume_indicators += 1
                
                # Check for email addresses (contact info)
                if re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text):
                    resume_indicators += 2
                
                # Check for dates (employment history)
                if re.search(r'\b\d{4}\b', resume_text):
                    resume_indicators += 1
                
                # Check for phone numbers
                if re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', resume_text):
                    resume_indicators += 1
                
                # Check for common resume sections
                resume_sections = ['summary', 'objective', 'profile', 'qualifications', 'achievements', 'certifications', 'awards', 'languages', 'references']
                for section in resume_sections:
                    if section.lower() in resume_text.lower():
                        resume_indicators += 1
                        break
                
                logger.info(f"Resume content indicators found: {resume_indicators}")
                logger.debug(f"Text preview (first 200 chars): {resume_text[:200]}")
                
                if resume_indicators < 1:  # Temporarily more lenient for debugging
                    logger.warning(f"Content doesn't look like a resume. Resume indicators: {resume_indicators}")
                    
                    # Provide helpful debugging info
                    text_preview = resume_text[:500] if len(resume_text) > 500 else resume_text
                    error_details = f"Text preview: {text_preview[:100]}..." if len(text_preview) > 100 else text_preview
                    logger.info(f"Rejected text details: {error_details}")
                    
                    return jsonify({
                        'success': False,
                        'error': 'The provided text does not appear to be resume content. Please ensure you are submitting actual resume text that includes work experience, education, skills, or contact information. If you uploaded a PDF file, try using the file upload button instead of copy/pasting.',
                        'help_tip': 'Click the file upload button and select your PDF/DOCX file, or copy clean text from your resume.',
                        'debug_info': {
                            'text_length': len(resume_text),
                            'resume_indicators': resume_indicators,
                            'text_preview': text_preview[:200] if len(text_preview) > 200 else text_preview
                        }
                    }), 400
                
                # Check for obvious binary content that wasn't properly extracted
                binary_indicators = 0
                
                # Check for PDF content indicators (be more strict)
                if resume_text.startswith('%PDF-') or '%PDF-' in resume_text[:100]:
                    logger.warning("PDF header detected in text - likely raw PDF content")
                    binary_indicators += 2  # Higher weight for PDF content
                    
                if resume_text.startswith('PK\x03\x04') or 'PK\x03\x04' in resume_text[:100]:
                    logger.warning("ZIP/DOCX signature detected - likely raw binary")
                    binary_indicators += 2  # Higher weight for binary headers
                    
                # Count null bytes (more strict)
                null_count = resume_text.count('\x00')
                if null_count > 5:  # Any significant null bytes indicate binary
                    logger.warning(f"Null bytes detected: {null_count}")
                    binary_indicators += 1
                    
                # Check for excessive non-printable characters (more strict)
                non_printable = [c for c in resume_text[:1000] if not c.isprintable() and c not in '\n\r\t\f\v']
                non_printable_ratio = len(non_printable) / len(resume_text[:1000]) if resume_text else 0
                if non_printable_ratio > 0.15:  # More than 15% non-printable
                    logger.warning(f"High non-printable character ratio: {non_printable_ratio:.2%}")
                    binary_indicators += 1
                
                # Check for common binary patterns
                binary_patterns = [
                    b'\x89PNG',  # PNG header
                    b'JFIF',     # JPEG header
                    b'\xff\xd8',  # JPEG start
                    b'%!PS',      # PostScript
                ]
                
                for pattern in binary_patterns:
                    if pattern.decode('latin-1', errors='ignore') in resume_text[:200]:
                        logger.warning(f"Binary pattern detected: {pattern}")
                        binary_indicators += 1
                        break
                
                # Check if text looks like base64 encoded content
                if len(resume_text) > 1000:
                    # Look for long strings of base64-like characters
                    base64_like = re.findall(r'[A-Za-z0-9+/]{50,}', resume_text)
                    if len(base64_like) > 5:
                        logger.warning("Possible base64 encoded content detected")
                        binary_indicators += 1
                
                logger.info(f"Binary indicators detected: {binary_indicators}")
                
                # Be more strict - flag as corrupted with fewer indicators
                if binary_indicators >= 2:
                    return jsonify({
                        'success': False,
                        'error': 'The content appears to be in PDF or binary format rather than readable text. Please: 1) Copy and paste the actual TEXT content from your resume (not the file), 2) Convert your PDF to text format first, or 3) Use a different file upload method.'
                    }), 400
                
                # Additional check for readability
                if len(resume_text) > 500:
                    # Sample first 500 characters and check readability
                    sample_text = resume_text[:500]
                    readable_chars = sum(1 for c in sample_text if c.isalnum() or c.isspace() or c in '.,!?-()[]{}:;"\'')
                    readability_ratio = readable_chars / len(sample_text)
                    
                    if readability_ratio < 0.7:  # Less than 70% readable characters
                        logger.warning(f"Low readability ratio: {readability_ratio:.2%}")
                        return jsonify({
                            'success': False,
                            'error': 'The text content appears to contain formatting or encoding issues. Please copy and paste clean, readable text from your resume, or try uploading a different file format.'
                        }), 400
                
                # Perform comprehensive analysis
                logger.info("Starting resume analysis...")
                analysis = self.perform_resume_analysis(resume_text, job_description, industry)
                logger.info(f"Analysis completed with overall score: {analysis.get('overall_score', 'unknown')}")
                
                # Save analysis to database
                session_id = session.get('session_id', self.generate_session_id())
                session['session_id'] = session_id
                
                self.save_analysis(session_id, resume_text, job_description, industry, analysis)
                
                # Track analytics
                self.track_event('resume_analysis_completed', {
                    'score': analysis['overall_score'],
                    'industry': industry,
                    'has_job_description': bool(job_description),
                    'file_upload': file_upload_attempted
                })
                
                return jsonify({
                    'success': True,
                    'analysis': analysis
                })
                
            except Exception as e:
                logger.error(f"Resume analysis error: {e}")
                return jsonify({
                    'success': False,
                    'error': 'Analysis failed due to an internal error. Please try again or contact support if the problem persists.'
                }), 500
        
        @self.app.route('/api/generate-cover-letter', methods=['POST'])
        def generate_cover_letter():
            """Generate personalized cover letter"""
            try:
                data = request.get_json()
                
                resume_text = data.get('resume_text', '')
                job_description = data.get('job_description', '')
                company_name = data.get('company_name', '')
                position_title = data.get('position_title', '')
                
                if not all([resume_text, company_name, position_title]):
                    return jsonify({
                        'success': False,
                        'error': 'Resume text, company name, and position title are required'
                    }), 400
                
                cover_letter = self.generate_cover_letter_content(
                    resume_text, job_description, company_name, position_title
                )
                
                # Track analytics
                self.track_event('cover_letter_generated', {
                    'company': company_name,
                    'position': position_title
                })
                
                return jsonify({
                    'success': True,
                    'cover_letter': cover_letter
                })
                
            except Exception as e:
                logger.error(f"Cover letter generation error: {e}")
                return jsonify({
                    'success': False,
                    'error': 'Cover letter generation failed. Please try again.'
                }), 500
        
        @self.app.route('/api/find-jobs', methods=['POST'])
        def find_jobs():
            """Find matching jobs based on resume"""
            try:
                data = request.get_json()
                
                resume_text = data.get('resume_text', '')
                location = data.get('location', 'United States')
                industry = data.get('industry', 'technology')
                
                if not resume_text:
                    return jsonify({
                        'success': False,
                        'error': 'Resume text is required'
                    }), 400
                
                jobs, search_mode = self.find_matching_jobs_with_mode(resume_text, location, industry)
                
                # Track analytics
                self.track_event('job_search_completed', {
                    'location': location,
                    'industry': industry,
                    'jobs_found': len(jobs),
                    'search_mode': search_mode
                })
                
                return jsonify({
                    'success': True,
                    'jobs': jobs,
                    'total_found': len(jobs),
                    'search_mode': search_mode
                })
                
            except Exception as e:
                logger.error(f"Job search endpoint error: {e}")
                return jsonify({
                    'success': False,
                    'error': 'Job search failed'
                }), 500
        
        @self.app.route('/api/job-search-status', methods=['GET'])
        def job_search_status():
            """Check job search API status"""
            try:
                # Check if we have API keys for real job search
                api_key = os.getenv('RAPIDAPI_KEY') or os.getenv('JSEARCH_API_KEY')
                has_requests = HAS_REQUESTS
                has_openai = bool(self.openai_client)
                
                real_jobs_available = bool(api_key and has_requests)
                curated_jobs_available = True  # Always available as fallback
                
                return jsonify({
                    'success': True,
                    'real_jobs_available': real_jobs_available,
                    'curated_jobs_available': curated_jobs_available,
                    'ai_enhancement_available': has_openai,
                    'api_configured': bool(api_key),
                    'requests_available': has_requests
                })
                
            except Exception as e:
                logger.error(f"Job search status check failed: {e}")
                return jsonify({
                    'success': False,
                    'real_jobs_available': False,
                    'curated_jobs_available': True,
                    'ai_enhancement_available': False
                })
                
            except Exception as e:
                logger.error(f"Job search error: {e}")
                return jsonify({
                    'success': False,
                    'error': 'Job search failed. Please try again.'
                }), 500
        
        @self.app.route('/api/interview-prep', methods=['POST'])
        def interview_prep():
            """Generate interview preparation materials"""
            try:
                data = request.get_json()
                
                resume_text = data.get('resume_text', '')
                job_description = data.get('job_description', '')
                interview_type = data.get('type', 'behavioral')
                
                if not resume_text:
                    return jsonify({
                        'success': False,
                        'error': 'Resume text is required'
                    }), 400
                
                prep_data = self.generate_interview_prep(resume_text, job_description, interview_type)
                
                # Track analytics
                self.track_event('interview_prep_generated', {
                    'type': interview_type
                })
                
                return jsonify({
                    'success': True,
                    'interview_prep': prep_data
                })
                
            except Exception as e:
                logger.error(f"Interview prep error: {e}")
                return jsonify({
                    'success': False,
                    'error': 'Interview preparation failed. Please try again.'
                }), 500
        
        @self.app.route('/api/grammar-check', methods=['POST'])
        def grammar_check():
            """Real-time grammar and spell checking"""
            try:
                data = request.get_json()
                text = data.get('text', '')
                
                if not text:
                    return jsonify({
                        'success': False,
                        'error': 'Text is required'
                    }), 400
                
                grammar_result = self.check_grammar_and_spelling(text)
                
                return jsonify({
                    'success': True,
                    'grammar_result': grammar_result
                })
                
            except Exception as e:
                logger.error(f"Grammar check error: {e}")
                return jsonify({
                    'success': False,
                    'error': 'Grammar check failed. Please try again.'
                }), 500
        
        @self.app.route('/api/tailor-resume', methods=['POST'])
        def tailor_resume():
            """Generate tailored resume with suggestions and truth-filtering"""
            try:
                data = request.get_json()
                
                resume_text = data.get('resume_text', '')
                job_description = data.get('job_description', '')
                analysis_context = data.get('analysis_context', {})
                
                if not resume_text:
                    return jsonify({
                        'success': False,
                        'error': 'Resume text is required'
                    }), 400
                
                if not job_description:
                    return jsonify({
                        'success': False,
                        'error': 'Job description is required'
                    }), 400
                
                # Generate tailored resume with truth-filtering and analysis context
                result = self.tailor_resume_with_context(resume_text, job_description, analysis_context)
                
                # Track analytics
                self.track_event('resume_tailored', {
                    'success': result['success'],
                    'changes_count': result.get('changes_count', 0),
                    'has_analysis_context': bool(analysis_context)
                })
                
                return jsonify(result)
                
            except Exception as e:
                logger.error(f"Resume tailoring error: {e}")
                return jsonify({
                    'success': False,
                    'error': 'Resume tailoring failed. Please try again.'
                }), 500
    
    def tailor_resume(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """
        Generate tailored resume suggestions with truth-filtering and JSON validation
        """
        return self.tailor_resume_with_context(resume_text, job_description, {})
    
    def tailor_resume_with_context(self, resume_text: str, job_description: str, analysis_context: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate tailored resume suggestions with analysis context and truth-filtering
        """
        if not self.openai_client:
            return {
                'success': False,
                'error': 'AI service not available',
                'original_sections': self.parse_resume_sections(resume_text),
                'suggestions': []
            }
        
        try:
            # Parse original resume into sections
            original_sections = self.parse_resume_sections(resume_text)
            
            # Generate tailored version with analysis context
            tailored_sections = self.generate_tailored_sections_with_context(
                resume_text, job_description, original_sections, analysis_context
            )
            
            # Apply truth filter
            validated_sections = self.apply_truth_filter(original_sections, tailored_sections)
            
            # Generate specific suggestions with analysis context
            suggestions = self.generate_improvement_suggestions_with_context(
                original_sections, validated_sections, job_description, analysis_context
            )
            
            return {
                'success': True,
                'original_sections': original_sections,
                'tailored_sections': validated_sections,
                'suggestions': suggestions,
                'changes_count': self.count_changes(original_sections, validated_sections),
                'analysis_insights': self.generate_analysis_insights(analysis_context)
            }
            
        except Exception as e:
            logger.error(f"Resume tailoring failed: {e}")
            return {
                'success': False,
                'error': str(e),
                'original_sections': self.parse_resume_sections(resume_text),
                'suggestions': []
            }
    
    def parse_resume_sections(self, resume_text: str) -> Dict[str, Any]:
        """Parse resume into structured sections"""
        sections = {
            'summary': '',
            'experience': [],
            'skills_core': [],
            'skills_plus': [],
            'education': [],
            'projects': []
        }
        
        lines = resume_text.split('\n')
        current_section = None
        current_item = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Detect section headers
            lower_line = line.lower()
            if any(header in lower_line for header in ['summary', 'objective', 'profile']):
                current_section = 'summary'
            elif any(header in lower_line for header in ['experience', 'employment', 'work history']):
                current_section = 'experience'
            elif any(header in lower_line for header in ['skills', 'technical skills', 'core competencies']):
                current_section = 'skills'
            elif any(header in lower_line for header in ['education', 'academic']):
                current_section = 'education'
            elif any(header in lower_line for header in ['projects', 'portfolio']):
                current_section = 'projects'
            else:
                # Process content based on current section
                if current_section == 'summary' and not any(char in line for char in ['•', '-', '◦']):
                    sections['summary'] += line + ' '
                elif current_section == 'experience':
                    # Simple experience parsing
                    if any(char in line for char in ['•', '-', '◦']):
                        if 'bullets' not in current_item:
                            current_item['bullets'] = []
                        current_item['bullets'].append(line.lstrip('•-◦ '))
                    elif '|' in line or any(year in line for year in ['2020', '2021', '2022', '2023', '2024', '2025']):
                        if current_item:
                            sections['experience'].append(current_item)
                        parts = line.split('|') if '|' in line else [line]
                        current_item = {
                            'title': parts[0].strip(),
                            'org': parts[1].strip() if len(parts) > 1 else 'Company',
                            'dates': parts[2].strip() if len(parts) > 2 else 'Recent',
                            'bullets': []
                        }
                elif current_section == 'skills':
                    # Extract skills
                    skills = [s.strip() for s in re.split(r'[,•\-◦]', line) if s.strip()]
                    sections['skills_core'].extend(skills[:3])  # First few as core
                    sections['skills_plus'].extend(skills[3:])  # Rest as plus
        
        # Add last experience item
        if current_item and current_section == 'experience':
            sections['experience'].append(current_item)
        
        # Clean up sections
        sections['summary'] = sections['summary'].strip()
        sections['skills_core'] = list(set(sections['skills_core']))[:8]
        sections['skills_plus'] = list(set(sections['skills_plus']))[:12]
        
        return sections
    
    def generate_tailored_sections(self, resume_text: str, job_description: str, original_sections: Dict[str, Any]) -> Dict[str, Any]:
        """Generate tailored resume sections using GPT with strict JSON output"""
        return self.generate_tailored_sections_with_context(resume_text, job_description, original_sections, {})
    
    def generate_tailored_sections_with_context(self, resume_text: str, job_description: str, original_sections: Dict[str, Any], analysis_context: Dict[str, Any]) -> Dict[str, Any]:
        """Generate tailored resume sections using GPT with analysis context"""
        
        # Build context-aware prompt
        context_info = ""
        if analysis_context:
            missing_keywords = analysis_context.get('missing_keywords', [])
            weak_areas = analysis_context.get('weak_areas', [])
            ats_score = analysis_context.get('ats_score', 0)
            
            if missing_keywords:
                context_info += f"\nMISSING KEYWORDS TO INCLUDE: {', '.join(missing_keywords[:10])}"
            if weak_areas:
                context_info += f"\nWEAK AREAS TO IMPROVE: {'; '.join(weak_areas[:5])}"
            if ats_score < 70:
                context_info += f"\nATS COMPATIBILITY: Currently {ats_score}% - needs improvement with better keyword integration"
        
        prompt = f"""You are a precise ATS résumé editor. Output strict JSON ONLY with this exact structure:
{{
  "summary_lines": ["line1", "line2", "line3"],
  "experience": [
    {{
      "title": "job title",
      "org": "company name", 
      "dates": "date range",
      "bullets": ["bullet1", "bullet2", "bullet3"]
    }}
  ],
  "skills_core": ["skill1", "skill2", "skill3"],
  "skills_plus": ["skill4", "skill5", "skill6"]
}}

RULES:
- NO new employers/companies not in original résumé
- NO new dates outside original date ranges  
- Prefer action verbs + quantified metrics
- MUST include missing keywords naturally where appropriate
- Focus on ATS optimization and keyword density
- Keep original employers and date ranges EXACTLY
- Strengthen weak areas identified in analysis

{context_info}

JOB DESCRIPTION:
{job_description[:1000]}

ORIGINAL RESUME:
{resume_text[:1500]}"""

        response = self.openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a precise résumé editor focused on ATS optimization. Output valid JSON only. No markdown, no explanations."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1200,
            temperature=0.3
        )
        
        content = response.choices[0].message.content.strip()
        
        # Clean and parse JSON
        if content.startswith('```json'):
            content = content[7:]
        if content.endswith('```'):
            content = content[:-3]
        
        try:
            return json.loads(content)
        except json.JSONDecodeError as e:
            logger.warning(f"JSON parsing failed, retrying: {e}")
            # Retry with more explicit instruction
            return self.retry_json_generation(resume_text, job_description)
    
    def retry_json_generation(self, resume_text: str, job_description: str) -> Dict[str, Any]:
        """Retry JSON generation with more explicit formatting"""
        prompt = f"""Generate ONLY valid JSON for résumé tailoring. No other text.

{{
  "summary_lines": ["Professional software engineer with expertise in...", "Experienced in full-stack development and...", "Proven track record of delivering..."],
  "experience": [
    {{
      "title": "Software Engineer",
      "org": "Tech Company",
      "dates": "2022-2024", 
      "bullets": ["Developed scalable applications...", "Improved system performance by...", "Collaborated with cross-functional teams..."]
    }}
  ],
  "skills_core": ["Python", "JavaScript", "React"],
  "skills_plus": ["Docker", "AWS", "Git"]
}}

Job Requirements: {job_description[:500]}
Original Resume: {resume_text[:800]}"""

        response = self.openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=800,
            temperature=0.1
        )
        
        try:
            return json.loads(response.choices[0].message.content.strip())
        except:
            # Final fallback
            return {
                "summary_lines": ["Experienced professional focused on delivering results"],
                "experience": [],
                "skills_core": [],
                "skills_plus": []
            }
    
    def apply_truth_filter(self, original_sections: Dict[str, Any], tailored_sections: Dict[str, Any]) -> Dict[str, Any]:
        """Apply truth filter to prevent hallucinated information"""
        validated = dict(tailored_sections)
        
        # Extract all employers and dates from original
        original_employers = set()
        original_dates = set()
        
        for exp in original_sections.get('experience', []):
            original_employers.add(exp.get('org', '').lower())
            original_dates.add(exp.get('dates', ''))
        
        # Filter experience entries
        validated_experience = []
        for exp in tailored_sections.get('experience', []):
            org_name = exp.get('org', '').lower()
            date_range = exp.get('dates', '')
            
            # Only include if employer exists in original
            if org_name in original_employers or any(org in org_name for org in original_employers):
                # Filter bullets for unseen technologies/companies
                filtered_bullets = []
                for bullet in exp.get('bullets', []):
                    if not self.contains_hallucinated_content(bullet, original_sections):
                        filtered_bullets.append(bullet)
                
                if filtered_bullets:  # Only include if we have valid bullets
                    exp['bullets'] = filtered_bullets
                    validated_experience.append(exp)
        
        validated['experience'] = validated_experience
        
        # Filter skills against original resume
        original_text = str(original_sections).lower()
        validated['skills_core'] = [skill for skill in tailored_sections.get('skills_core', []) 
                                  if skill.lower() in original_text or self.is_reasonable_skill_inference(skill, original_text)]
        validated['skills_plus'] = [skill for skill in tailored_sections.get('skills_plus', []) 
                                  if skill.lower() in original_text or self.is_reasonable_skill_inference(skill, original_text)]
        
        return validated
    
    def contains_hallucinated_content(self, bullet: str, original_sections: Dict[str, Any]) -> bool:
        """Check if bullet contains information not in original resume"""
        bullet_lower = bullet.lower()
        original_text = str(original_sections).lower()
        
        # Check for company names not in original
        company_indicators = ['at ', 'for ', 'with ', ' inc', ' corp', ' ltd', ' llc']
        for indicator in company_indicators:
            if indicator in bullet_lower:
                # Extract potential company name
                start = bullet_lower.find(indicator)
                potential_company = bullet_lower[start:start+30].split()[1:3]
                if potential_company and not any(word in original_text for word in potential_company):
                    return True
        
        return False
    
    def is_reasonable_skill_inference(self, skill: str, original_text: str) -> bool:
        """Check if skill is a reasonable inference from original resume"""
        skill_lower = skill.lower()
        
        # Allow common related skills
        skill_families = {
            'javascript': ['js', 'react', 'node', 'angular', 'vue'],
            'python': ['django', 'flask', 'pandas', 'numpy'],
            'aws': ['cloud', 'amazon', 'ec2', 's3'],
            'docker': ['container', 'kubernetes', 'k8s'],
            'git': ['github', 'version control', 'gitlab']
        }
        
        for family, related in skill_families.items():
            if skill_lower == family:
                return any(term in original_text for term in related)
            if skill_lower in related:
                return family in original_text
        
        return False
    
    def generate_improvement_suggestions(self, original: Dict[str, Any], tailored: Dict[str, Any], job_description: str) -> List[Dict[str, Any]]:
        """Generate specific improvement suggestions with before/after examples"""
        return self.generate_improvement_suggestions_with_context(original, tailored, job_description, {})
    
    def generate_improvement_suggestions_with_context(self, original: Dict[str, Any], tailored: Dict[str, Any], job_description: str, analysis_context: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate specific improvement suggestions with analysis context"""
        suggestions = []
        
        # Add analysis-driven suggestions first
        if analysis_context:
            missing_keywords = analysis_context.get('missing_keywords', [])
            ats_score = analysis_context.get('ats_score', 0)
            action_verb_score = analysis_context.get('action_verb_score', 0)
            
            # Missing keywords suggestions
            if missing_keywords:
                suggestions.append({
                    'type': 'keyword_optimization',
                    'section': 'Overall Resume',
                    'before': 'Current keyword coverage',
                    'after': f'Add these keywords naturally: {", ".join(missing_keywords[:5])}',
                    'reason': f'Missing {len(missing_keywords)} important keywords that appear in the job description',
                    'priority': 'high'
                })
            
            # ATS optimization suggestions
            if ats_score < 70:
                suggestions.append({
                    'type': 'ats_optimization',
                    'section': 'ATS Compatibility',
                    'before': f'Current ATS score: {ats_score}%',
                    'after': 'Improved keyword density and ATS-friendly formatting',
                    'reason': 'Low ATS compatibility score - resume may not pass automated screening',
                    'priority': 'high'
                })
            
            # Action verb suggestions
            if action_verb_score < 60:
                suggestions.append({
                    'type': 'action_verb_improvement',
                    'section': 'Experience Bullets',
                    'before': 'Weak action verbs and passive language',
                    'after': 'Strong action verbs: Developed, Implemented, Achieved, Led, Optimized',
                    'reason': 'Improve impact with stronger action verbs',
                    'priority': 'medium'
                })
        
        # Compare experience bullets
        for i, orig_exp in enumerate(original.get('experience', [])):
            if i < len(tailored.get('experience', [])):
                tailored_exp = tailored['experience'][i]
                
                # Compare bullets
                orig_bullets = orig_exp.get('bullets', [])
                new_bullets = tailored_exp.get('bullets', [])
                
                for j, orig_bullet in enumerate(orig_bullets):
                    if j < len(new_bullets) and orig_bullet != new_bullets[j]:
                        suggestions.append({
                            'type': 'bullet_improvement',
                            'section': f'Experience - {orig_exp.get("title", "Role")}',
                            'before': orig_bullet,
                            'after': new_bullets[j],
                            'reason': 'Better alignment with job requirements and stronger action verbs',
                            'priority': 'medium'
                        })
        
        # Compare summary
        orig_summary = original.get('summary', '')
        new_summary = ' '.join(tailored.get('summary_lines', []))
        if orig_summary and new_summary and orig_summary != new_summary:
            suggestions.append({
                'type': 'summary_improvement',
                'section': 'Professional Summary',
                'before': orig_summary,
                'after': new_summary,
                'reason': 'More targeted to job requirements and includes relevant keywords',
                'priority': 'high'
            })
        
        # Skills additions
        orig_skills = set(original.get('skills_core', []) + original.get('skills_plus', []))
        new_skills = set(tailored.get('skills_core', []) + tailored.get('skills_plus', []))
        added_skills = new_skills - orig_skills
        
        if added_skills:
            suggestions.append({
                'type': 'skills_addition',
                'section': 'Technical Skills',
                'before': 'Current skills list',
                'after': f'Add: {", ".join(added_skills)}',
                'reason': 'Highlighted skills relevant to this role from job description',
                'priority': 'medium'
            })
        
        # Sort suggestions by priority
        priority_order = {'high': 0, 'medium': 1, 'low': 2}
        suggestions.sort(key=lambda x: priority_order.get(x.get('priority', 'low'), 2))
        
        return suggestions
    
    def generate_analysis_insights(self, analysis_context: Dict[str, Any]) -> Dict[str, Any]:
        """Generate insights from analysis context"""
        if not analysis_context:
            return {}
        
        insights = {
            'keyword_gaps': len(analysis_context.get('missing_keywords', [])),
            'ats_readiness': analysis_context.get('ats_score', 0) >= 70,
            'action_verb_strength': analysis_context.get('action_verb_score', 0) >= 60,
            'improvement_areas': []
        }
        
        # Identify improvement areas
        if insights['keyword_gaps'] > 5:
            insights['improvement_areas'].append('Keyword optimization')
        if not insights['ats_readiness']:
            insights['improvement_areas'].append('ATS compatibility')
        if not insights['action_verb_strength']:
            insights['improvement_areas'].append('Action verb strength')
        
        return insights
    
    def count_changes(self, original: Dict[str, Any], tailored: Dict[str, Any]) -> int:
        """Count the number of changes made"""
        changes = 0
        
        # Count bullet changes
        for i, orig_exp in enumerate(original.get('experience', [])):
            if i < len(tailored.get('experience', [])):
                tailored_exp = tailored['experience'][i]
                orig_bullets = orig_exp.get('bullets', [])
                new_bullets = tailored_exp.get('bullets', [])
                
                for j, orig_bullet in enumerate(orig_bullets):
                    if j < len(new_bullets) and orig_bullet != new_bullets[j]:
                        changes += 1
        
        # Count summary changes
        if original.get('summary') != ' '.join(tailored.get('summary_lines', [])):
            changes += 1
        
        # Count skill changes
        orig_skills = set(original.get('skills_core', []) + original.get('skills_plus', []))
        new_skills = set(tailored.get('skills_core', []) + tailored.get('skills_plus', []))
        if orig_skills != new_skills:
            changes += 1
        
        return changes

    def extract_text_from_uploaded_file(self, file):
        """Extract text from uploaded file (PDF, DOCX, TXT)"""
        try:
            if not file or not file.filename:
                logger.warning("No file or filename provided")
                return ""
                
            filename = file.filename.lower()
            logger.info(f"Extracting text from file: {filename}")
            
            # Reset file pointer to beginning
            file.seek(0)
            
            if filename.endswith('.pdf'):
                return self.extract_pdf_text(file)
                
            elif filename.endswith('.docx'):
                return self.extract_docx_text(file)
                
            elif filename.endswith('.txt'):
                return self.extract_txt_text(file)
                
            else:
                # Try to read as text file (fallback)
                return self.extract_fallback_text(file)
                
        except Exception as e:
            logger.error(f"General file extraction error: {e}")
            return ""
    
    def extract_pdf_text(self, file):
        """Extract text from PDF using available libraries"""
        text = ""
        
        try:
            # Try pdfplumber first (better text extraction)
            if HAS_PDFPLUMBER:
                logger.info("Attempting PDF extraction with pdfplumber")
                file.seek(0)
                with pdfplumber.open(file) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                                logger.debug(f"Extracted {len(page_text)} chars from page {page_num + 1}")
                        except Exception as page_error:
                            logger.warning(f"Failed to extract text from page {page_num + 1}: {page_error}")
                
                if text.strip():
                    logger.info(f"Successfully extracted {len(text)} characters from PDF using pdfplumber")
                    return text.strip()
        
        except Exception as pdfplumber_error:
            logger.warning(f"pdfplumber extraction failed: {pdfplumber_error}")
        
        # Fallback to PyPDF2
        try:
            if HAS_PDF:
                logger.info("Attempting PDF extraction with PyPDF2")
                file.seek(0)
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            logger.debug(f"PyPDF2 extracted {len(page_text)} chars from page {page_num + 1}")
                    except Exception as page_error:
                        logger.warning(f"PyPDF2 failed on page {page_num + 1}: {page_error}")
                
                if text.strip():
                    logger.info(f"Successfully extracted {len(text)} characters from PDF using PyPDF2")
                    return text.strip()
        
        except Exception as pypdf2_error:
            logger.warning(f"PyPDF2 extraction failed: {pypdf2_error}")
        
        logger.warning("All PDF extraction methods failed")
        return ""
    
    def extract_docx_text(self, file):
        """Extract text from DOCX file"""
        try:
            if not HAS_DOCX:
                logger.warning("python-docx not available")
                return ""
            
            file.seek(0)
            doc = Document(file)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted {len(text)} characters from DOCX")
                return text.strip()
            else:
                logger.warning("DOCX extraction returned empty text")
                return ""
                
        except Exception as docx_error:
            logger.error(f"DOCX extraction error: {docx_error}")
            return ""
    
    def extract_txt_text(self, file):
        """Extract text from TXT file"""
        try:
            file.seek(0)
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    text = file.read().decode(encoding).strip()
                    if text:
                        logger.info(f"Successfully extracted {len(text)} characters from TXT using {encoding}")
                        return text
                except UnicodeDecodeError:
                    continue
            
            logger.warning("Could not decode TXT file with any encoding")
            return ""
            
        except Exception as txt_error:
            logger.error(f"TXT extraction error: {txt_error}")
            return ""
    
    def extract_fallback_text(self, file):
        """Fallback text extraction for unknown formats"""
        try:
            file.seek(0)
            # Try to read as text with error handling
            content = file.read()
            
            # Try different decodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    text = content.decode(encoding, errors='ignore').strip()
                    if text and len(text) > 10:
                        # Check if it looks like readable text (not binary)
                        readable_chars = sum(1 for c in text[:500] if c.isprintable() or c.isspace())
                        if readable_chars / len(text[:500]) > 0.7:  # 70% readable characters
                            logger.info(f"Successfully extracted {len(text)} characters from unknown format using {encoding}")
                            return text
                except:
                    continue
            
            logger.warning("Fallback text extraction failed - file may be binary")
            return ""
            
        except Exception as fallback_error:
            logger.error(f"Fallback extraction error: {fallback_error}")
            return ""
            return ""
    
    def perform_resume_analysis(self, resume_text: str, job_description: str, industry: str) -> Dict[str, Any]:
        """
        Comprehensive resume analysis using multiple AI techniques with content-specific scoring
        """
        # Create content hash for unique analysis caching
        import hashlib
        content_hash = hashlib.md5(f"{resume_text[:500]}{job_description[:200]}{industry}".encode()).hexdigest()[:8]
        
        analysis = {
            'overall_score': 0,
            'grammar_score': 0,
            'ats_score': 0,
            'action_verb_score': 0,
            'quantified_score': 0,
            'cliche_score': 0,
            'matched_keywords': [],
            'missing_keywords': [],
            'suggestions': [],
            'detailed_feedback': {},
            'content_hash': content_hash
        }
        
        # Use AI for comprehensive analysis if available
        ai_analysis_success = False
        if self.openai_client:
            try:
                ai_analysis = self.get_ai_powered_analysis(resume_text, job_description, industry)
                if ai_analysis and ai_analysis.get('overall_score', 0) > 0:
                    # Merge AI analysis with traditional analysis
                    analysis.update(ai_analysis)
                    ai_analysis_success = True
                    logger.info(f"AI analysis successful for content hash: {content_hash}")
                else:
                    logger.warning("AI analysis returned empty or invalid results")
            except Exception as e:
                logger.warning(f"AI analysis failed, using traditional methods: {e}")
        
        # Enhanced traditional analysis (with content-specific variations)
        # 1. Grammar and Spelling Analysis
        grammar_result = self.check_grammar_and_spelling(resume_text)
        traditional_grammar_score = max(0, 100 - (grammar_result['error_count'] * 5))
        
        # 2. ATS Compatibility Analysis
        traditional_ats_score = self.analyze_ats_compatibility(resume_text)
        
        # 3. Action Verb Analysis
        traditional_action_score = self.analyze_action_verbs(resume_text)
        
        # 4. Quantified Achievements Analysis
        traditional_quantified_score = self.analyze_quantified_achievements(resume_text)
        
        # 5. Cliché Detection
        traditional_cliche_score = self.detect_cliches(resume_text)
        
        # Use AI scores if available, otherwise use traditional scores
        if not ai_analysis_success:
            analysis['grammar_score'] = traditional_grammar_score
            analysis['ats_score'] = traditional_ats_score
            analysis['action_verb_score'] = traditional_action_score
            analysis['quantified_score'] = traditional_quantified_score
            analysis['cliche_score'] = traditional_cliche_score
            logger.info(f"Using traditional analysis for content hash: {content_hash}")
        else:
            # Blend AI and traditional scores for more accuracy
            analysis['grammar_score'] = int((analysis.get('grammar_score', 70) + traditional_grammar_score) / 2)
            analysis['ats_score'] = int((analysis.get('ats_score', 70) + traditional_ats_score) / 2)
            analysis['action_verb_score'] = int((analysis.get('action_verb_score', 70) + traditional_action_score) / 2)
            analysis['quantified_score'] = int((analysis.get('quantified_score', 70) + traditional_quantified_score) / 2)
            analysis['cliche_score'] = int((analysis.get('cliche_score', 70) + traditional_cliche_score) / 2)
        
        # 6. Keyword Matching (if job description provided)
        if job_description:
            keyword_analysis = self.analyze_keywords(resume_text, job_description)
            analysis['matched_keywords'] = keyword_analysis['matched']
            analysis['missing_keywords'] = keyword_analysis['missing']
        
        # 7. Generate AI-powered suggestions
        try:
            analysis['suggestions'] = self.generate_suggestions(resume_text, job_description, industry)
        except:
            analysis['suggestions'] = ["Focus on quantifying achievements", "Use stronger action verbs", "Ensure ATS compatibility"]
        
        # Calculate overall score (content-specific)
        if 'overall_score' not in analysis or analysis['overall_score'] == 0:
            scores = [
                analysis['grammar_score'],
                analysis['ats_score'],
                analysis['action_verb_score'],
                analysis['quantified_score'],
                analysis['cliche_score']
            ]
            
            # Apply content-based adjustments to prevent identical scores
            base_score = int(sum(scores) / len(scores))
            
            # Content-specific adjustments
            resume_length = len(resume_text.split())
            if resume_length < 200:
                base_score -= 5  # Too short
            elif resume_length > 800:
                base_score += 3   # Comprehensive
            
            # Industry bonus/penalty
            industry_keywords = {
                'technology': ['software', 'code', 'development', 'programming', 'technical'],
                'healthcare': ['patient', 'medical', 'clinical', 'healthcare', 'treatment'],
                'finance': ['financial', 'analysis', 'budget', 'accounting', 'investment'],
                'marketing': ['campaign', 'brand', 'digital', 'social media', 'analytics']
            }
            
            if industry in industry_keywords:
                relevant_terms = sum(1 for term in industry_keywords[industry] 
                                   if term.lower() in resume_text.lower())
                if relevant_terms >= 3:
                    base_score += 2
                elif relevant_terms == 0:
                    base_score -= 3
            
            analysis['overall_score'] = max(0, min(100, base_score))
        
        return analysis
    
    def get_ai_powered_analysis(self, resume_text: str, job_description: str, industry: str) -> Dict[str, Any]:
        """
        Get AI-powered comprehensive resume analysis with dynamic scoring
        """
        # Create a unique analysis prompt that varies based on content
        resume_length = len(resume_text)
        has_job_desc = bool(job_description.strip())
        word_count = len(resume_text.split())
        
        prompt = f"""As an expert resume analyst, provide a comprehensive, personalized analysis of this resume. Consider the specific content, industry context, and job requirements to give unique, tailored feedback.

RESUME TO ANALYZE ({word_count} words):
{resume_text[:2500]}

{"JOB DESCRIPTION TO MATCH:" if has_job_desc else "GENERAL INDUSTRY ANALYSIS for:"}
{(job_description[:1000] if has_job_desc else f"Industry: {industry} - Provide general industry-specific recommendations")}

ANALYSIS REQUIREMENTS:
1. Provide specific, content-based scoring (not generic)
2. Consider the actual skills, experience, and achievements mentioned
3. Factor in resume length, structure, and presentation quality
4. Give industry-specific feedback for {industry}
5. Identify unique strengths and specific improvement areas

Score each area 0-100 based on actual content quality:
- Overall Quality: Holistic assessment of this specific resume
- Grammar/Professionalism: Based on actual writing quality found
- ATS Compatibility: Keyword usage, formatting, structure for this resume
- Action Verbs: Strength and variety of verbs actually used
- Quantified Results: Presence and quality of metrics in this resume
- Uniqueness: How well this resume stands out from generic templates

Provide detailed, specific feedback about:
- What makes this resume strong (specific examples)
- Concrete areas for improvement (with examples)
- Missing elements that would help this candidate
- Industry-specific recommendations for {industry}

Return ONLY valid JSON in this exact format:
{{
    "overall_score": 75,
    "grammar_score": 88,
    "ats_score": 82,
    "action_verb_score": 67,
    "quantified_score": 71,
    "cliche_score": 79,
    "strengths": ["Specific strength 1", "Specific strength 2", "Specific strength 3"],
    "improvements": ["Specific improvement 1", "Specific improvement 2", "Specific improvement 3"],
    "missing_elements": ["Missing element 1", "Missing element 2"],
    "industry_feedback": "Industry-specific feedback for {industry}",
    "uniqueness_score": 76
}}"""

        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": f"You are a senior resume analyst with 20+ years of experience in {industry} recruiting. Provide honest, specific, and actionable feedback. Each resume is unique - analyze the actual content, not generic patterns. Vary your scoring based on what you actually observe in the resume content."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.4  # Slightly higher for more varied responses
            )
            
            # Try to parse JSON response
            ai_response = response.choices[0].message.content.strip()
            logger.info(f"AI Response received: {len(ai_response)} characters")
            
            # Extract JSON from response (sometimes AI includes extra text)
            json_start = ai_response.find('{')
            json_end = ai_response.rfind('}') + 1
            
            if json_start >= 0 and json_end > json_start:
                json_str = ai_response[json_start:json_end]
                ai_analysis = json.loads(json_str)
                
                # Validate scores are reasonable and different
                required_keys = ['overall_score', 'grammar_score', 'ats_score', 'action_verb_score', 'quantified_score', 'cliche_score']
                for key in required_keys:
                    if key not in ai_analysis:
                        ai_analysis[key] = 70  # Default fallback
                    # Ensure scores are in valid range
                    ai_analysis[key] = max(0, min(100, int(ai_analysis[key])))
                
                logger.info(f"AI Analysis successful - Overall Score: {ai_analysis.get('overall_score', 'unknown')}")
                return ai_analysis
            else:
                logger.warning("Could not parse AI analysis JSON - no valid JSON found")
                return {}
                
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse AI analysis JSON: {e}")
            logger.debug(f"Raw AI response: {ai_response}")
            return {}
        except Exception as e:
            logger.warning(f"AI analysis failed: {e}")
            return {}
    
    def check_grammar_and_spelling(self, text: str) -> Dict[str, Any]:
        """
        Grammar and spelling check with improved detection
        """
        error_count = 0
        errors = []
        suggestions = []
        
        # Common spelling errors
        spelling_errors = {
            r'\brecieve\b': 'receive',
            r'\bseperate\b': 'separate', 
            r'\bdefinately\b': 'definitely',
            r'\boccured\b': 'occurred',
            r'\baccomodate\b': 'accommodate',
            r'\bexperiance\b': 'experience',
            r'\bmanagment\b': 'management',
            r'\bsuccessfull\b': 'successful',
            r'\bresponsibile\b': 'responsible',
            r'\bperfomance\b': 'performance'
        }
        
        for error_pattern, correction in spelling_errors.items():
            matches = re.findall(error_pattern, text, re.IGNORECASE)
            if matches:
                error_count += len(matches)
                errors.append(f"Spelling: '{matches[0]}' should be '{correction}'")
        
        # Grammar issues
        grammar_patterns = [
            (r'\.{2,}', "Multiple periods found - use single period"),
            (r',{2,}', "Multiple commas found - use single comma"),
            (r'\s{3,}', "Multiple spaces found - use single space"),
            (r'\b(I|i)\s+are\b', "Subject-verb disagreement: 'I are' should be 'I am'"),
            (r'\b(He|She|It)\s+are\b', "Subject-verb disagreement: use 'is' instead of 'are'"),
            (r'\bwho\s+are\s+responsible\b', "Consider: 'who is responsible' or 'who are responsible for'"),
            (r'\bthere\s+resume\b', "Should be 'their resume'"),
            (r'\bits\s+own\b', "Consider: 'its own' (no apostrophe for possessive 'its')"),
        ]
        
        for pattern, message in grammar_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                error_count += 1
                errors.append(f"Grammar: {message}")
        
        # Style suggestions
        style_issues = [
            (r'\bvery\s+\w+', "Consider stronger adjectives instead of 'very + adjective'"),
            (r'\ba\s+lot\s+of\b', "Consider 'many' or 'numerous' instead of 'a lot of'"),
            (r'\bbasically\b', "Avoid filler words like 'basically'"),
            (r'\bobviously\b', "Avoid assumptive words like 'obviously'"),
        ]
        
        for pattern, suggestion in style_issues:
            if re.search(pattern, text, re.IGNORECASE):
                suggestions.append(suggestion)
        
        # Check for inconsistent formatting
        if re.search(r'\d{4}\s*-\s*\d{4}', text) and re.search(r'\d{4}\s*/\s*\d{4}', text):
            errors.append("Inconsistent date formatting (mix of dashes and slashes)")
            error_count += 1
        
        return {
            'error_count': error_count,
            'errors': errors[:10],  # Limit to first 10 errors
            'suggestions': suggestions[:5]  # Limit to first 5 suggestions
        }
    
    def analyze_ats_compatibility(self, resume_text: str) -> int:
        """
        Analyze ATS (Applicant Tracking System) compatibility with improved scoring
        """
        score = 100
        text_length = len(resume_text)
        
        # Check for problematic formatting (more nuanced)
        special_chars = len(re.findall(r'[^\w\s\-\.\,\(\)\[\]\/\@\#\%\&\*\+\=\:\;\<\>\?\!]', resume_text))
        if special_chars > text_length * 0.05:  # More than 5% special characters
            score -= min(20, special_chars // 10)
        
        # Check for standard sections with better detection
        sections_found = 0
        section_patterns = {
            'experience': r'(experience|employment|work\s+history|professional\s+background)',
            'education': r'(education|academic|university|college|degree)',
            'skills': r'(skills|competencies|technologies|technical|core\s+skills)',
            'contact': r'(@|email|phone|\d{3}[-.]?\d{3}[-.]?\d{4})'
        }
        
        for section, pattern in section_patterns.items():
            if re.search(pattern, resume_text, re.IGNORECASE):
                sections_found += 1
            else:
                score -= 12
        
        # Bonus for having all sections
        if sections_found == len(section_patterns):
            score += 10
        
        # Check for dates with better patterns
        date_patterns = [
            r'\d{4}',  # Year
            r'\d{1,2}/\d{4}',  # Month/Year
            r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{4}',  # Month Year
            r'(present|current)',  # Current employment
        ]
        dates_found = sum(len(re.findall(pattern, resume_text, re.IGNORECASE)) for pattern in date_patterns)
        if dates_found == 0:
            score -= 15
        elif dates_found < 2:
            score -= 8
        
        # Check for contact information
        if not re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text):
            score -= 15
        
        # Check for phone number
        if not re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', resume_text):
            score -= 5
        
        # Length considerations
        word_count = len(resume_text.split())
        if word_count < 200:
            score -= 20  # Too short
        elif word_count > 1000:
            score -= 5   # Might be too long
        
        return max(0, min(100, score))
    
    def analyze_action_verbs(self, resume_text: str) -> int:
        """
        Analyze the use of strong action verbs with improved scoring
        """
        strong_verbs = [
            'achieved', 'improved', 'increased', 'developed', 'implemented',
            'managed', 'led', 'created', 'designed', 'optimized', 'delivered',
            'reduced', 'enhanced', 'streamlined', 'initiated', 'spearheaded',
            'built', 'launched', 'established', 'executed', 'coordinated',
            'supervised', 'facilitated', 'analyzed', 'transformed', 'generated',
            'collaborated', 'mentored', 'negotiated', 'resolved', 'exceeded'
        ]
        
        weak_verbs = ['responsible for', 'duties included', 'worked on', 'helped with', 'was involved']
        
        # Count strong verbs (case insensitive)
        strong_verb_count = 0
        for verb in strong_verbs:
            strong_verb_count += len(re.findall(r'\b' + verb + r'\w*', resume_text, re.IGNORECASE))
        
        # Count weak phrases and penalize
        weak_phrase_count = 0
        for phrase in weak_verbs:
            weak_phrase_count += len(re.findall(phrase, resume_text, re.IGNORECASE))
        
        # Count total sentences/bullet points for context
        sentences = len(re.findall(r'[.!?]|\n\s*[•\-\*]', resume_text))
        if sentences == 0:
            sentences = len(resume_text.split('\n'))
        
        # Calculate density of strong action verbs
        if sentences == 0:
            return 0
        
        strong_density = strong_verb_count / max(1, sentences)
        weak_penalty = weak_phrase_count * 5
        
        # Scoring based on density and variety
        if strong_density >= 0.8:  # 80% of sentences have strong verbs
            score = 95
        elif strong_density >= 0.6:
            score = 85
        elif strong_density >= 0.4:
            score = 70
        elif strong_density >= 0.2:
            score = 55
        else:
            score = 30
        
        # Variety bonus - check for diverse verb usage
        unique_verbs_used = set()
        for verb in strong_verbs:
            if re.search(r'\b' + verb + r'\w*', resume_text, re.IGNORECASE):
                unique_verbs_used.add(verb)
        
        variety_bonus = min(15, len(unique_verbs_used) * 2)
        score += variety_bonus
        
        # Apply weak phrase penalty
        score = max(20, score - weak_penalty)
        
        return min(100, score)
    
    def analyze_quantified_achievements(self, resume_text: str) -> int:
        """
        Analyze the presence of quantified achievements with better detection
        """
        # More comprehensive patterns for quantification
        patterns = [
            r'\d+%',  # percentages (15%, 25%)
            r'\$\d+[,\d]*\.?\d*[kmb]?',  # dollar amounts ($50K, $1.2M)
            r'\d+[,\d]*\+?\s*(million|billion|thousand|k\b)',  # large numbers
            r'\d+[,\d]*\s*(years?|months?|weeks?|days?)',  # time periods
            r'\d+[,\d]*\s*(people|employees|team\s+members|clients|customers)',  # team/client sizes
            r'\d+[,\d]*\s*(projects?|applications?|systems?|processes?)',  # work quantities
            r'(increased|decreased|improved|reduced|grew|cut|saved)\s+\w*\s*by\s+\d+',  # improvement metrics
            r'\d+[,\d]*x\s',  # multiplier (10x faster)
            r'(top|#)\s*\d+',  # rankings (top 5)
            r'\d+[,\d]*\s*(hours?|minutes?)\s*(per|\/)',  # time savings
            r'over\s+\d+[,\d]*',  # over X amount
            r'up\s+to\s+\d+[,\d]*',  # up to X amount
            r'\d+[,\d]*\s*-\s*\d+[,\d]*',  # ranges (10-15)
        ]
        
        quantified_count = 0
        achievements_found = []
        
        for pattern in patterns:
            matches = re.findall(pattern, resume_text, re.IGNORECASE)
            quantified_count += len(matches)
            achievements_found.extend(matches)
        
        # Look for achievement context words near numbers
        achievement_contexts = [
            'revenue', 'sales', 'profit', 'efficiency', 'performance', 'productivity',
            'cost', 'budget', 'roi', 'conversion', 'growth', 'retention',
            'satisfaction', 'quality', 'compliance', 'accuracy', 'speed'
        ]
        
        context_bonus = 0
        for context in achievement_contexts:
            # Look for numbers within 20 characters of achievement words
            context_pattern = rf'\b{context}\b.{{0,20}}\d+|\d+.{{0,20}}\b{context}\b'
            context_matches = len(re.findall(context_pattern, resume_text, re.IGNORECASE))
            context_bonus += context_matches * 2
        
        # Calculate base score
        if quantified_count >= 12:
            base_score = 95
        elif quantified_count >= 8:
            base_score = 85
        elif quantified_count >= 5:
            base_score = 70
        elif quantified_count >= 3:
            base_score = 55
        elif quantified_count >= 1:
            base_score = 40
        else:
            base_score = 15
        
        # Add context bonus (up to 15 points)
        total_score = min(100, base_score + min(15, context_bonus))
        
        return total_score
    
    def detect_cliches(self, resume_text: str) -> int:
        """
        Detect and penalize clichéd phrases
        """
        cliches = [
            'team player', 'hard worker', 'detail oriented', 'self motivated',
            'results driven', 'excellent communication skills', 'fast learner',
            'go getter', 'out of the box', 'hit the ground running'
        ]
        
        cliche_count = 0
        for cliche in cliches:
            cliche_count += len(re.findall(cliche, resume_text, re.IGNORECASE))
        
        # Higher cliché count = lower score
        return max(20, 100 - (cliche_count * 15))
    
    def analyze_keywords(self, resume_text: str, job_description: str) -> Dict[str, List[str]]:
        """
        Analyze keyword matching between resume and job description
        """
        if not job_description:
            return {'matched': [], 'missing': []}
        
        # Extract keywords from job description
        job_keywords = self.extract_keywords(job_description)
        resume_keywords = self.extract_keywords(resume_text)
        
        matched = list(set(job_keywords) & set(resume_keywords))
        missing = list(set(job_keywords) - set(resume_keywords))
        
        return {
            'matched': matched[:10],  # Limit to top 10
            'missing': missing[:10]   # Limit to top 10
        }
    
    def extract_keywords(self, text: str) -> List[str]:
        """
        Extract important keywords from text
        """
        # Simple keyword extraction (in production, use more sophisticated NLP)
        
        # Common technical skills and buzzwords
        keywords = []
        
        # Programming languages
        prog_langs = ['python', 'java', 'javascript', 'react', 'node.js', 'sql', 'html', 'css']
        for lang in prog_langs:
            if re.search(r'\b' + lang + r'\b', text, re.IGNORECASE):
                keywords.append(lang)
        
        # Skills
        skills = ['leadership', 'management', 'analysis', 'communication', 'problem solving']
        for skill in skills:
            if re.search(r'\b' + skill + r'\b', text, re.IGNORECASE):
                keywords.append(skill)
        
        # Extract capitalized words (likely to be important)
        cap_words = re.findall(r'\b[A-Z][a-z]+\b', text)
        keywords.extend(cap_words[:5])  # Take first 5
        
        return list(set(keywords))
    
    def generate_suggestions(self, resume_text: str, job_description: str, industry: str) -> List[str]:
        """
        Generate AI-powered improvement suggestions
        """
        # Use OpenAI GPT for advanced suggestions if available
        if self.openai_client:
            try:
                gpt_suggestions = self.generate_gpt_suggestions(resume_text, job_description, industry)
                if gpt_suggestions:
                    return gpt_suggestions
            except Exception as e:
                logger.warning(f"GPT suggestions generation failed, using fallback: {e}")
        
        # Fallback to rule-based suggestions
        return self.generate_rule_based_suggestions(resume_text, job_description, industry)
    
    def generate_gpt_suggestions(self, resume_text: str, job_description: str, industry: str) -> List[str]:
        """
        Generate AI-powered suggestions using GPT-4
        """
        prompt = f"""Analyze this resume and provide 6-8 specific, actionable improvement suggestions:

Resume:
{resume_text[:2000]}

Job Description (if provided):
{job_description[:1000] if job_description else 'No specific job description provided'}

Industry: {industry}

Provide specific, actionable suggestions to improve this resume. Focus on:
1. Content improvements (missing skills, experiences to highlight)
2. Formatting and structure improvements
3. Keyword optimization for ATS systems
4. Quantification opportunities
5. Industry-specific recommendations
6. Professional language enhancements

Return exactly 6-8 bullet points, each being a specific actionable suggestion."""

        response = self.openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert resume coach with 15+ years of experience helping professionals get hired at top companies. Provide specific, actionable advice."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=600,
            temperature=0.7
        )
        
        suggestions_text = response.choices[0].message.content.strip()
        
        # Parse the suggestions from the response
        suggestions = []
        for line in suggestions_text.split('\n'):
            line = line.strip()
            if line and (line.startswith('•') or line.startswith('-') or line.startswith('*')):
                # Remove bullet point markers
                suggestion = line.lstrip('•-*').strip()
                if suggestion:
                    suggestions.append(suggestion)
            elif line and not line.startswith(('Here', 'Based on', 'To improve')):
                # Catch suggestions that don't have bullet points
                suggestions.append(line)
        
        return suggestions[:8]  # Limit to 8 suggestions
    
    def generate_rule_based_suggestions(self, resume_text: str, job_description: str, industry: str) -> List[str]:
        """
        Generate rule-based suggestions (fallback)
        """
        suggestions = []
        
        # Basic suggestions based on analysis
        if len(resume_text.split()) < 200:
            suggestions.append("Consider expanding your resume with more detailed descriptions of your experiences.")
        
        if not re.search(r'\d+%', resume_text):
            suggestions.append("Add quantified achievements with percentages to demonstrate your impact.")
        
        if not re.search(r'\$\d+', resume_text):
            suggestions.append("Include monetary values where possible to show the financial impact of your work.")
        
        if job_description and len(self.analyze_keywords(resume_text, job_description)['missing']) > 5:
            suggestions.append("Consider incorporating more keywords from the job description into your resume.")
        
        if not re.search(r'\b(led|managed|supervised)\b', resume_text, re.IGNORECASE):
            suggestions.append("Highlight leadership experiences using strong action verbs like 'led', 'managed', or 'supervised'.")
        
        # Industry-specific suggestions
        if industry == 'technology':
            if not re.search(r'\b(github|portfolio|projects)\b', resume_text, re.IGNORECASE):
                suggestions.append("Consider adding links to your GitHub profile or portfolio to showcase your technical projects.")
        
        elif industry == 'marketing':
            if not re.search(r'\b(campaign|roi|conversion)\b', resume_text, re.IGNORECASE):
                suggestions.append("Include metrics from marketing campaigns such as ROI, conversion rates, or engagement metrics.")
        
        return suggestions[:8]  # Limit to 8 suggestions
    
    def generate_cover_letter_content(self, resume_text: str, job_description: str, company_name: str, position_title: str) -> str:
        """
        Generate personalized cover letter using AI or template
        """
        # Use OpenAI GPT if available
        if self.openai_client:
            try:
                return self.generate_gpt_cover_letter(resume_text, job_description, company_name, position_title)
            except Exception as e:
                logger.warning(f"GPT cover letter generation failed, falling back to template: {e}")
        
        # Fallback to template-based approach
        return self.generate_template_cover_letter(resume_text, job_description, company_name, position_title)
    
    def generate_gpt_cover_letter(self, resume_text: str, job_description: str, company_name: str, position_title: str) -> str:
        """
        Generate cover letter using OpenAI GPT-4
        """
        prompt = f"""Write a professional, personalized cover letter for the following job application:

Position: {position_title}
Company: {company_name}

Job Description:
{job_description[:1500]}

Candidate's Resume:
{resume_text[:2000]}

Instructions:
- Write a compelling cover letter that highlights relevant experience from the resume
- Match the tone to the job description
- Include specific examples from the candidate's background
- Keep it professional but engaging
- Length: 3-4 paragraphs
- Do not include placeholder text like [Your Name] - write as if the candidate is speaking
- Focus on value proposition and cultural fit

Cover Letter:"""

        response = self.openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert career coach and professional writer who specializes in creating compelling cover letters that get interviews."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=800,
            temperature=0.7
        )
        
        return response.choices[0].message.content.strip()
    
    def generate_template_cover_letter(self, resume_text: str, job_description: str, company_name: str, position_title: str) -> str:
        """
        Generate template-based cover letter (fallback)
        """
        # Extract key skills from resume
        skills = self.extract_keywords(resume_text)[:3]
        
        # Generate template-based cover letter
        cover_letter = f"""Dear Hiring Manager,

I am writing to express my strong interest in the {position_title} position at {company_name}. With my background in {', '.join(skills)} and proven track record of delivering results, I am confident I would be a valuable addition to your team.

In my previous roles, I have demonstrated expertise in {skills[0] if skills else 'relevant technologies'} and have consistently delivered high-quality solutions. My experience has taught me the importance of collaboration, innovation, and continuous learning – values that align perfectly with {company_name}'s mission.

Key highlights from my background include:
• Strong technical skills in {skills[0] if skills else 'relevant technologies'}
• Experience with {skills[1] if len(skills) > 1 else 'project management'}
• Proven ability to {skills[2] if len(skills) > 2 else 'deliver results'}

I am particularly excited about the opportunity to contribute to {company_name}'s continued success and would welcome the chance to discuss how my skills and experience can benefit your team.

Thank you for your time and consideration. I look forward to hearing from you.

Sincerely,
[Your Name]"""
        
        return cover_letter
    
    def find_matching_jobs(self, resume_text: str, location: str, industry: str) -> List[Dict[str, Any]]:
        """
        Find matching jobs using real job search APIs
        """
        jobs, _ = self.find_matching_jobs_with_mode(resume_text, location, industry)
        return jobs
    
    def find_matching_jobs_with_mode(self, resume_text: str, location: str, industry: str) -> tuple[List[Dict[str, Any]], str]:
        """
        Find matching jobs with search mode information
        Returns: (jobs_list, search_mode)
        """
        try:
            # Try real job search first
            jobs = self.search_real_jobs(resume_text, location, industry)
            
            if jobs:
                search_mode = 'real_jobs'
                # Enhance with AI if available
                if self.openai_client:
                    jobs = self.enhance_jobs_with_ai_matching(jobs, resume_text, industry)
                else:
                    jobs = self.apply_traditional_matching(jobs, resume_text)
            else:
                # Fallback to curated jobs if API fails
                jobs = self.get_curated_jobs(industry, location)
                search_mode = 'curated'
                if self.openai_client:
                    jobs = self.enhance_jobs_with_ai_matching(jobs, resume_text, industry)
            
            # Sort by match score and return top results
            jobs.sort(key=lambda x: x.get('match_score', 0), reverse=True)
            return jobs[:8], search_mode  # Return top 8 matches
            
        except Exception as e:
            logger.error(f"Job search failed: {e}")
            fallback_jobs = self.get_fallback_jobs(industry, location)
            return fallback_jobs, 'fallback'
    
    def search_real_jobs(self, resume_text: str, location: str, industry: str) -> List[Dict[str, Any]]:
        """
        Search for real jobs using multiple APIs with intelligent fallback
        """
        if not HAS_REQUESTS:
            logger.warning("Requests library not available, using fallback jobs")
            return []
            
        # Try LinkedIn first (if available)
        linkedin_jobs = self.search_linkedin_jobs(resume_text, location, industry)
        if linkedin_jobs:
            logger.info(f"Found {len(linkedin_jobs)} LinkedIn jobs")
            return linkedin_jobs
        
        # Fallback to JSearch API
        jsearch_jobs = self.search_jsearch_jobs(resume_text, location, industry)
        if jsearch_jobs:
            logger.info(f"Found {len(jsearch_jobs)} JSearch jobs")
            return jsearch_jobs
            
        logger.warning("No real jobs found from APIs")
        return []
    
    def search_linkedin_jobs(self, resume_text: str, location: str, industry: str) -> List[Dict[str, Any]]:
        """
        Search for jobs using LinkedIn Job Search API (RapidAPI)
        """
        try:
            # Extract key skills from resume for better search
            key_skills = self.extract_key_skills_from_resume(resume_text)
            
            # LinkedIn Job Search API configuration
            url = "https://linkedin-job-search-api.p.rapidapi.com/active-jb-1h"
            
            # Get API key from environment
            api_key = os.getenv('RAPIDAPI_KEY') or os.getenv('JSEARCH_API_KEY')
            
            if not api_key:
                logger.warning("No RapidAPI key found for LinkedIn search")
                return []
            
            headers = {
                "X-RapidAPI-Key": api_key,
                "X-RapidAPI-Host": "linkedin-job-search-api.p.rapidapi.com"
            }
            
            # Search parameters - get fresh jobs from last hour
            params = {
                "offset": "0"  # Start from the beginning
            }
            
            # Make API request with timeout
            response = requests.get(url, headers=headers, params=params, timeout=15)
            
            if response.status_code == 200:
                data = response.json()
                
                # Handle different possible response structures
                jobs_data = []
                if isinstance(data, list):
                    jobs_data = data
                elif isinstance(data, dict):
                    jobs_data = data.get('jobs', data.get('data', [data]))
                
                # Filter jobs by industry and location preferences
                filtered_jobs = self.filter_linkedin_jobs(jobs_data, location, industry, key_skills)
                
                # Convert API response to our format
                formatted_jobs = []
                for job in filtered_jobs[:12]:  # Get up to 12 jobs for variety
                    formatted_job = self.format_linkedin_job(job)
                    if formatted_job:
                        formatted_jobs.append(formatted_job)
                
                logger.info(f"LinkedIn API success: {len(formatted_jobs)} jobs formatted")
                return formatted_jobs
            
            elif response.status_code == 401:
                logger.warning("LinkedIn API: Subscription or endpoint access restricted")
                return []
            else:
                logger.warning(f"LinkedIn Job API returned status {response.status_code}")
                return []
                
        except requests.RequestException as e:
            logger.warning(f"LinkedIn Job API request failed: {e}")
            return []
        except Exception as e:
            logger.error(f"LinkedIn job search failed: {e}")
            return []
    
    def search_jsearch_jobs(self, resume_text: str, location: str, industry: str) -> List[Dict[str, Any]]:
        """
        Search for jobs using JSearch API (RapidAPI) - Fallback option
        """
        try:
            # Extract key skills from resume for better search
            key_skills = self.extract_key_skills_from_resume(resume_text)
            
            # Build search query
            search_query = self.build_job_search_query(industry, key_skills)
            
            # JSearch API configuration
            url = "https://jsearch.p.rapidapi.com/search"
            
            # Get API key from environment
            api_key = os.getenv('RAPIDAPI_KEY') or os.getenv('JSEARCH_API_KEY')
            
            if not api_key:
                logger.warning("No RapidAPI key found for JSearch")
                return []
            
            headers = {
                "X-RapidAPI-Key": api_key,
                "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
            }
            
            # Search parameters
            params = {
                "query": search_query,
                "page": "1",
                "num_pages": "1",
                "date_posted": "week",  # Jobs from last week
                "employment_types": "FULLTIME,PARTTIME,CONTRACTOR",
                "job_requirements": "under_3_years_experience,more_than_3_years_experience"
            }
            
            # Add location if provided
            if location and location.strip() and location.lower() != "united states":
                params["location"] = location.strip()
            
            # Make API request with timeout
            response = requests.get(url, headers=headers, params=params, timeout=15)
            
            if response.status_code == 200:
                data = response.json()
                jobs = data.get('data', [])
                
                # Convert API response to our format
                formatted_jobs = []
                for job in jobs[:10]:  # Get up to 10 jobs for variety
                    formatted_job = self.format_jsearch_job(job)
                    if formatted_job:
                        formatted_jobs.append(formatted_job)
                
                logger.info(f"JSearch API success: {len(formatted_jobs)} jobs formatted from query: {search_query}")
                return formatted_jobs
            
            else:
                logger.warning(f"JSearch API returned status {response.status_code}")
                return []
                
        except requests.RequestException as e:
            logger.warning(f"JSearch API request failed: {e}")
            return []
        except Exception as e:
            logger.error(f"JSearch job search failed: {e}")
            return []
    
    def extract_key_skills_from_resume(self, resume_text: str) -> List[str]:
        """
        Extract key skills from resume for targeted job search
        """
        # Common tech skills and keywords
        tech_skills = [
            'python', 'javascript', 'java', 'c++', 'react', 'angular', 'vue',
            'node.js', 'django', 'flask', 'spring', 'sql', 'mysql', 'postgresql',
            'mongodb', 'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'git',
            'machine learning', 'data science', 'ai', 'tensorflow', 'pytorch',
            'html', 'css', 'typescript', 'php', 'ruby', 'go', 'rust', 'scala'
        ]
        
        # Business skills
        business_skills = [
            'project management', 'agile', 'scrum', 'marketing', 'sales',
            'customer service', 'business analysis', 'consulting', 'strategy',
            'operations', 'finance', 'accounting', 'hr', 'recruiting'
        ]
        
        resume_lower = resume_text.lower()
        found_skills = []
        
        # Find tech skills
        for skill in tech_skills:
            if skill in resume_lower:
                found_skills.append(skill)
        
        # Find business skills
        for skill in business_skills:
            if skill in resume_lower:
                found_skills.append(skill)
        
        # Extract job titles/roles mentioned
        title_patterns = [
            r'software engineer', r'data scientist', r'product manager',
            r'marketing manager', r'sales manager', r'business analyst',
            r'project manager', r'developer', r'programmer', r'consultant'
        ]
        
        for pattern in title_patterns:
            if re.search(pattern, resume_lower):
                found_skills.append(pattern.replace(r'\b', '').replace(r'\s+', ' '))
        
        return found_skills[:5]  # Return top 5 skills
    
    def filter_linkedin_jobs(self, jobs_data: List[Dict], location: str, industry: str, key_skills: List[str]) -> List[Dict]:
        """
        Filter LinkedIn jobs by relevance to user preferences
        """
        if not jobs_data:
            return []
        
        filtered_jobs = []
        location_lower = location.lower() if location else ""
        industry_lower = industry.lower() if industry else ""
        skills_lower = [skill.lower() for skill in key_skills]
        
        for job in jobs_data:
            try:
                # Get job details
                title = job.get('title', '').lower()
                company = job.get('company', '').lower()
                description = job.get('description', '').lower()
                job_location = job.get('location', '').lower()
                
                # Calculate relevance score
                relevance_score = 0
                
                # Industry matching
                if industry_lower:
                    if industry_lower in title or industry_lower in description:
                        relevance_score += 30
                    elif industry_lower in ['tech', 'technology'] and any(tech_word in title + description for tech_word in ['software', 'developer', 'engineer', 'tech', 'programming']):
                        relevance_score += 25
                    elif industry_lower in ['marketing'] and any(marketing_word in title + description for marketing_word in ['marketing', 'brand', 'campaign', 'digital']):
                        relevance_score += 25
                
                # Skills matching
                for skill in skills_lower:
                    if skill in title:
                        relevance_score += 15
                    elif skill in description:
                        relevance_score += 10
                
                # Location preference (less strict for remote)
                if location_lower and location_lower != "united states":
                    if location_lower in job_location:
                        relevance_score += 20
                    elif 'remote' in job_location:
                        relevance_score += 15
                
                # Include jobs with reasonable relevance
                if relevance_score >= 15 or not industry_lower:  # Include all if no specific industry
                    job['relevance_score'] = relevance_score
                    filtered_jobs.append(job)
                    
            except Exception as e:
                logger.warning(f"Error filtering job: {e}")
                continue
        
        # Sort by relevance score
        filtered_jobs.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
        return filtered_jobs
    
    def format_linkedin_job(self, job_data: Dict) -> Optional[Dict[str, Any]]:
        """
        Format LinkedIn job data to our standard format
        """
        try:
            # Extract key information with fallbacks
            title = job_data.get('title', 'Position Available')
            company = job_data.get('company', 'Company')
            location = job_data.get('location', 'Remote')
            description = job_data.get('description', 'Job description not available')
            
            # Clean up description
            if len(description) > 500:
                description = description[:500] + "..."
            
            # Extract/generate salary info
            salary = self.extract_linkedin_salary(job_data)
            
            # Get job URL
            job_url = job_data.get('url') or job_data.get('link') or job_data.get('job_url', '#')
            
            # Calculate base match score (enhanced by relevance if available)
            base_score = 75  # Default for LinkedIn jobs
            if 'relevance_score' in job_data:
                base_score = min(95, 60 + job_data['relevance_score'])
            
            # Get posting date
            posted_date = self.format_linkedin_date(job_data.get('posted_date') or job_data.get('date'))
            
            # Determine remote friendliness
            remote_friendly = bool(
                'remote' in location.lower() or 
                'remote' in description.lower() or 
                'work from home' in description.lower()
            )
            
            # Extract company size (if available)
            company_size = job_data.get('company_size', 'Not specified')
            
            # Extract benefits from description
            benefits = self.extract_benefits_from_description(description)
            
            # Add LinkedIn-specific benefits if not found
            if not benefits:
                benefits = ['Professional Development', 'LinkedIn Learning', 'Networking Opportunities']
            
            formatted_job = {
                "title": title,
                "company": company,
                "location": location,
                "salary": salary,
                "match_score": base_score,
                "description": description,
                "url": job_url,
                "posted_date": posted_date,
                "remote_friendly": remote_friendly,
                "company_size": company_size,
                "benefits": benefits,
                "source": "LinkedIn Jobs"
            }
            
            return formatted_job
            
        except Exception as e:
            logger.warning(f"Failed to format LinkedIn job data: {e}")
            return None
    
    def extract_linkedin_salary(self, job_data: Dict) -> str:
        """
        Extract and format salary information from LinkedIn job data
        """
        # Try different salary fields
        salary = job_data.get('salary')
        salary_min = job_data.get('salary_min')
        salary_max = job_data.get('salary_max')
        
        if salary:
            return str(salary)
        elif salary_min and salary_max:
            return f"${salary_min:,} - ${salary_max:,} per year"
        elif salary_min:
            return f"${salary_min:,}+ per year"
        else:
            # Generate competitive salary based on job title
            title_lower = job_data.get('title', '').lower()
            if 'senior' in title_lower or 'lead' in title_lower:
                return "$90,000 - $150,000 per year"
            elif 'manager' in title_lower or 'director' in title_lower:
                return "$100,000 - $180,000 per year"
            elif 'entry' in title_lower or 'junior' in title_lower:
                return "$50,000 - $80,000 per year"
            else:
                return "Competitive salary"
    
    def format_linkedin_date(self, date_str: str) -> str:
        """
        Format LinkedIn posting date to relative time
        """
        if not date_str:
            return "Recent"
        
        try:
            # Handle different date formats from LinkedIn
            date_str = str(date_str).lower()
            
            if 'hour' in date_str:
                return "Today"
            elif 'day' in date_str:
                # Extract number of days
                import re
                days_match = re.search(r'(\d+)', date_str)
                if days_match:
                    days = int(days_match.group(1))
                    if days == 1:
                        return "1 day ago"
                    else:
                        return f"{days} days ago"
                return "Recent"
            elif 'week' in date_str:
                weeks_match = re.search(r'(\d+)', date_str)
                if weeks_match:
                    weeks = int(weeks_match.group(1))
                    return f"{weeks} week{'s' if weeks > 1 else ''} ago"
                return "1 week ago"
            elif 'month' in date_str:
                return "1 month ago"
            else:
                return "Recent"
                
        except Exception as e:
            logger.warning(f"Error formatting date {date_str}: {e}")
            return "Recent"
    
    def build_job_search_query(self, industry: str, skills: List[str]) -> str:
        """
        Build optimized search query for job APIs
        """
        query_parts = []
        
        # Add industry
        if industry and industry.strip():
            query_parts.append(industry.strip())
        
        # Add top skills
        if skills:
            # Add most relevant skills
            query_parts.extend(skills[:3])
        
        # Default fallback
        if not query_parts:
            query_parts = ['software engineer', 'developer']
        
        return ' '.join(query_parts)
    
    def format_jsearch_job(self, job_data: Dict) -> Optional[Dict[str, Any]]:
        """
        Format JSearch job data to our standard format
        """
        try:
            # Extract key information with fallbacks
            title = job_data.get('job_title', 'Position Available')
            company = job_data.get('employer_name', 'Company')
            
            # Handle location
            city = job_data.get('job_city', '')
            state = job_data.get('job_state', '')
            country = job_data.get('job_country', 'Remote')
            
            if city and state:
                location = f"{city}, {state}"
            elif city:
                location = city
            elif state:
                location = state
            else:
                location = country
            
            description = job_data.get('job_description', 'Job description not available')
            
            # Clean up description
            if len(description) > 500:
                description = description[:500] + "..."
            
            # Extract salary info
            salary = self.extract_jsearch_salary(job_data)
            
            # Format job URL
            job_url = job_data.get('job_apply_link') or job_data.get('job_google_link', '#')
            
            # Calculate base match score
            base_score = 75  # Default for JSearch jobs
            
            # Get posting date
            posted_date = self.format_jsearch_date(job_data.get('job_posted_at_datetime_utc'))
            
            # Determine remote friendliness
            remote_friendly = bool(
                job_data.get('job_is_remote') or 
                'remote' in description.lower() or 
                'work from home' in description.lower()
            )
            
            # Extract company size (if available)
            company_size = "Not specified"
            
            # Extract benefits from description
            benefits = self.extract_benefits_from_description(description)
            
            # Add default benefits if none found
            if not benefits:
                benefits = ['Competitive Package', 'Professional Growth']
            
            formatted_job = {
                "title": title,
                "company": company,
                "location": location,
                "salary": salary,
                "match_score": base_score,
                "description": description,
                "url": job_url,
                "posted_date": posted_date,
                "remote_friendly": remote_friendly,
                "company_size": company_size,
                "benefits": benefits,
                "source": "Real Job Board"
            }
            
            return formatted_job
            
        except Exception as e:
            logger.warning(f"Failed to format JSearch job data: {e}")
            return None
    
    def extract_jsearch_salary(self, job_data: Dict) -> str:
        """
        Extract and format salary information from JSearch job data
        """
        # Try different salary fields
        min_salary = job_data.get('job_min_salary')
        max_salary = job_data.get('job_max_salary')
        salary_currency = job_data.get('job_salary_currency', 'USD')
        salary_period = job_data.get('job_salary_period', 'YEAR')
        
        if min_salary and max_salary:
            if salary_period == 'YEAR':
                return f"${min_salary:,} - ${max_salary:,} per year"
            elif salary_period == 'MONTH':
                return f"${min_salary:,} - ${max_salary:,} per month"
            elif salary_period == 'HOUR':
                return f"${min_salary} - ${max_salary} per hour"
            else:
                return f"${min_salary:,} - ${max_salary:,}"
        elif min_salary:
            return f"${min_salary:,}+ per year"
        else:
            return "Competitive salary"
    
    def format_jsearch_date(self, date_str: str) -> str:
        """
        Format JSearch posting date to relative time
        """
        if not date_str:
            return "Recent"
        
        try:
            # Convert ISO datetime to relative time
            from datetime import datetime
            posted_dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            now = datetime.now(posted_dt.tzinfo)
            days_ago = (now - posted_dt).days
            
            if days_ago == 0:
                return "Today"
            elif days_ago == 1:
                return "1 day ago"
            elif days_ago < 7:
                return f"{days_ago} days ago"
            elif days_ago < 30:
                weeks = days_ago // 7
                return f"{weeks} week{'s' if weeks > 1 else ''} ago"
            else:
                return "1 month ago"
                
        except Exception as e:
            logger.warning(f"Error formatting JSearch date {date_str}: {e}")
            return "Recent"
    
    def extract_benefits_from_description(self, description: str) -> List[str]:
        """
        Extract benefits from job description
        """
        benefits = []
        description_lower = description.lower()
        
        benefit_keywords = {
            "health insurance": ["health insurance", "medical insurance", "healthcare"],
            "401k": ["401k", "retirement", "pension"],
            "pto": ["pto", "paid time off", "vacation", "paid leave"],
            "remote work": ["remote work", "work from home", "telecommute"],
            "flexible hours": ["flexible hours", "flexible schedule", "flex time"],
            "stock options": ["stock options", "equity", "shares"],
            "professional development": ["professional development", "training", "learning", "education"],
            "bonus": ["bonus", "incentive"]
        }
        
        for benefit, keywords in benefit_keywords.items():
            if any(keyword in description_lower for keyword in keywords):
                benefits.append(benefit.title())
        
        return benefits[:4]  # Return top 4 benefits
    
    def get_curated_jobs(self, industry: str, location: str) -> List[Dict[str, Any]]:
        """
        Get curated job listings when API is unavailable
        """
    def get_curated_jobs(self, industry: str, location: str) -> List[Dict[str, Any]]:
        """
        Get curated job listings when API is unavailable
        """
        base_jobs = []
        
        # Technology jobs
        if industry.lower() in ['technology', 'tech', 'software', 'it', 'engineering']:
            base_jobs = [
                {
                    "title": "Senior Software Engineer",
                    "company": "Tech Innovation Corp",
                    "location": location or "San Francisco, CA",
                    "salary": "$120,000 - $160,000",
                    "match_score": 85,
                    "description": "Join our team to build scalable web applications using React, Node.js, and cloud technologies. We're looking for experienced developers who can architect robust solutions and mentor junior team members.",
                    "url": "https://careers.techinnovation.com/senior-engineer",
                    "posted_date": "2 days ago",
                    "remote_friendly": True,
                    "company_size": "500-1000",
                    "benefits": ["Health Insurance", "401k", "Flexible PTO", "Remote Work"],
                    "source": "Curated"
                },
                {
                    "title": "Full Stack Developer",
                    "company": "StartupXYZ",
                    "location": location or "New York, NY",
                    "salary": "$90,000 - $130,000",
                    "match_score": 80,
                    "description": "We're building the next generation of fintech solutions. Looking for a passionate developer proficient in Python/Django and React to help us scale our platform.",
                    "url": "https://startupxyz.com/careers/fullstack",
                    "posted_date": "1 day ago",
                    "remote_friendly": True,
                    "company_size": "50-100",
                    "benefits": ["Equity", "Health Insurance", "Professional Development"],
                    "source": "Curated"
                },
                {
                    "title": "DevOps Engineer",
                    "company": "Cloud Solutions Inc",
                    "location": location or "Austin, TX",
                    "salary": "$110,000 - $145,000",
                    "match_score": 75,
                    "description": "Lead infrastructure automation and CI/CD pipeline development. Experience with AWS, Docker, Kubernetes, and Infrastructure as Code required.",
                    "url": "https://cloudsolutions.com/jobs/devops",
                    "posted_date": "3 days ago",
                    "remote_friendly": False,
                    "company_size": "200-500",
                    "benefits": ["Health Insurance", "401k", "Professional Development"],
                    "source": "Curated"
                },
                {
                    "title": "Data Scientist",
                    "company": "Analytics Pro",
                    "location": location or "Seattle, WA",
                    "salary": "$95,000 - $135,000",
                    "match_score": 70,
                    "description": "Apply machine learning and statistical analysis to solve complex business problems. Python, SQL, and ML frameworks experience required.",
                    "url": "https://analyticspro.com/careers/data-scientist",
                    "posted_date": "1 week ago",
                    "remote_friendly": True,
                    "company_size": "100-200",
                    "benefits": ["Health Insurance", "Stock Options", "Flexible Hours"],
                    "source": "Curated"
                }
            ]
        
        # Marketing jobs
        elif industry.lower() in ['marketing', 'digital marketing', 'advertising']:
            base_jobs = [
                {
                    "title": "Digital Marketing Manager",
                    "company": "Growth Marketing Co",
                    "location": location or "Los Angeles, CA",
                    "salary": "$75,000 - $95,000",
                    "match_score": 85,
                    "description": "Lead digital marketing campaigns across multiple channels. Experience with Google Ads, Facebook Ads, and marketing automation required.",
                    "url": "https://growthmarketing.com/jobs/digital-manager",
                    "posted_date": "2 days ago",
                    "remote_friendly": True,
                    "company_size": "100-200",
                    "benefits": ["Health Insurance", "401k", "Professional Development"],
                    "source": "Curated"
                },
                {
                    "title": "Content Marketing Specialist",
                    "company": "Content Creators LLC",
                    "location": location or "Chicago, IL",
                    "salary": "$55,000 - $75,000",
                    "match_score": 80,
                    "description": "Create engaging content for blog, social media, and email campaigns. Strong writing skills and SEO knowledge required.",
                    "url": "https://contentcreators.com/careers/content-specialist",
                    "posted_date": "1 day ago",
                    "remote_friendly": True,
                    "company_size": "20-50",
                    "benefits": ["Health Insurance", "Flexible PTO", "Professional Development"],
                    "source": "Curated"
                }
            ]
        
        # Finance jobs
        elif industry.lower() in ['finance', 'banking', 'financial services']:
            base_jobs = [
                {
                    "title": "Financial Analyst",
                    "company": "Investment Partners LLC",
                    "location": location or "New York, NY",
                    "salary": "$70,000 - $95,000",
                    "match_score": 85,
                    "description": "Analyze financial data, prepare reports, and support investment decisions. Strong Excel and analytical skills required.",
                    "url": "https://investmentpartners.com/careers/analyst",
                    "posted_date": "1 day ago",
                    "remote_friendly": False,
                    "company_size": "200-500",
                    "benefits": ["Health Insurance", "401k", "Bonus"],
                    "source": "Curated"
                }
            ]
        
        # Default/other industries
        else:
            base_jobs = [
                {
                    "title": f"{industry.title()} Specialist",
                    "company": "Professional Services Corp",
                    "location": location or "Multiple Locations",
                    "salary": "$60,000 - $85,000",
                    "match_score": 75,
                    "description": f"Seeking experienced {industry} professional to join our growing team. Competitive benefits and growth opportunities available.",
                    "url": "https://professionalservices.com/careers",
                    "posted_date": "1 week ago",
                    "remote_friendly": False,
                    "company_size": "500+",
                    "benefits": ["Health Insurance", "401k", "PTO"],
                    "source": "Curated"
                }
            ]
        
        return base_jobs
    
    def enhance_jobs_with_ai_matching(self, jobs: List[Dict], resume_text: str, industry: str) -> List[Dict]:
        """
        Enhance job matching using AI analysis (optimized for speed)
        """
        try:
            # Batch process for efficiency
            for job in jobs:
                # Calculate AI-powered match score
                enhanced_score = self.calculate_ai_match_score(job, resume_text)
                job['match_score'] = min(100, max(job.get('match_score', 50), enhanced_score))
                
                # Add AI insights
                job['ai_insights'] = self.get_quick_ai_insights(job, resume_text)
            
            return jobs
            
        except Exception as e:
            logger.warning(f"AI job enhancement failed: {e}")
            return jobs
    
    def calculate_ai_match_score(self, job: Dict, resume_text: str) -> int:
        """
        Calculate AI-powered job match score (optimized)
        """
        try:
            # Quick AI analysis with reduced token usage
            prompt = f"""Rate job match (0-100) for this candidate:

Job: {job.get('title', '')} at {job.get('company', '')}
Resume: {resume_text[:800]}

Consider: skills match, experience level, title relevance.
Return only a number 0-100."""

            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a job matching expert. Provide accurate match scores."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=10,
                temperature=0.1
            )
            
            score_text = response.choices[0].message.content.strip()
            score = int(''.join(filter(str.isdigit, score_text)))
            return min(100, max(0, score))
            
        except Exception as e:
            logger.warning(f"AI match scoring failed: {e}")
            return 50
    
    def get_quick_ai_insights(self, job: Dict, resume_text: str) -> str:
        """
        Get quick AI insights about job match
        """
        try:
            prompt = f"""Brief insight on why this job matches the candidate:

Job: {job.get('title', '')}
Candidate skills: {resume_text[:400]}

One sentence insight on the match."""

            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Provide brief, helpful job match insights."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=50,
                temperature=0.3
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            logger.warning(f"AI insights failed: {e}")
            return "Good potential match based on your background."
    
    def apply_traditional_matching(self, jobs: List[Dict], resume_text: str) -> List[Dict]:
        """
        Apply traditional keyword-based matching (fallback)
        """
        resume_lower = resume_text.lower()
        
        for job in jobs:
            title_lower = job.get('title', '').lower()
            description_lower = job.get('description', '').lower()
            
            # Adjust scores based on keyword matching
            base_score = job.get('match_score', 50)
            
            # Title matching
            if any(word in resume_lower for word in title_lower.split()):
                base_score += 15
            
            # Skill matching
            skills = ['python', 'javascript', 'react', 'sql', 'aws', 'docker', 'kubernetes']
            matched_skills = sum(1 for skill in skills if skill in resume_lower and skill in description_lower)
            base_score += matched_skills * 5
            
            # Experience level matching
            if 'senior' in title_lower and 'senior' in resume_lower:
                base_score += 10
            if 'manager' in title_lower and ('manage' in resume_lower or 'lead' in resume_lower):
                base_score += 10
            
            job['match_score'] = min(100, base_score)
        
        return jobs
    
    def get_fallback_jobs(self, industry: str, location: str) -> List[Dict]:
        """
        Fallback job results when all search methods fail
        """
        return [
            {
                "title": f"{industry.title()} Opportunity",
                "company": "Various Companies",
                "location": location or "Multiple Locations",
                "salary": "Competitive",
                "match_score": 60,
                "description": f"Multiple {industry} opportunities available across different companies. Our job search system is being optimized to provide better real-time results.",
                "url": "#",
                "posted_date": "Recent",
                "remote_friendly": True,
                "company_size": "Various",
                "benefits": ["Competitive packages available"],
                "source": "Fallback"
            }
        ]
    
    def generate_interview_prep(self, resume_text: str, job_description: str, interview_type: str) -> Dict[str, Any]:
        """
        Generate interview preparation materials using AI
        """
        # Use OpenAI GPT for personalized interview prep if available
        if self.openai_client:
            try:
                return self.generate_gpt_interview_prep(resume_text, job_description, interview_type)
            except Exception as e:
                logger.warning(f"GPT interview prep generation failed, using fallback: {e}")
        
        # Fallback to template-based prep
        return self.generate_template_interview_prep(resume_text, job_description, interview_type)
    
    def generate_gpt_interview_prep(self, resume_text: str, job_description: str, interview_type: str) -> Dict[str, Any]:
        """
        Generate personalized interview prep using GPT-4
        """
        prompt = f"""Create personalized interview preparation for this candidate:

Interview Type: {interview_type}

Candidate's Resume:
{resume_text[:1500]}

Job Description:
{job_description[:1000] if job_description else 'General interview preparation'}

Generate:
1. 5-7 relevant interview questions for this specific candidate and role
2. 5-6 practical preparation tips tailored to their background
3. 2-3 specific examples they could use from their resume

Focus on {interview_type} interview style. Make it personal and specific to their experience."""

        response = self.openai_client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": f"You are an expert interview coach specializing in {interview_type} interviews. Provide specific, actionable advice tailored to the candidate's background."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        content = response.choices[0].message.content.strip()
        
        # Parse the response to extract questions and tips
        lines = content.split('\n')
        questions = []
        tips = []
        examples = []
        
        current_section = None
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if 'question' in line.lower() and ('1.' in line or ':' in line):
                current_section = 'questions'
                continue
            elif 'tip' in line.lower() and ('1.' in line or ':' in line):
                current_section = 'tips'
                continue
            elif 'example' in line.lower() and ('1.' in line or ':' in line):
                current_section = 'examples'
                continue
            
            if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '•', '-', '*')):
                clean_line = re.sub(r'^[\d\.\-\*\•]+\s*', '', line)
                if current_section == 'questions':
                    questions.append(clean_line)
                elif current_section == 'tips':
                    tips.append(clean_line)
                elif current_section == 'examples':
                    examples.append(clean_line)
        
        # Fallback parsing if structure isn't clear
        if not questions:
            questions = self.get_fallback_questions(interview_type)
        if not tips:
            tips = self.get_fallback_tips(interview_type)
        
        return {
            'questions': questions[:7],
            'tips': tips[:6],
            'examples': examples[:3] if examples else []
        }
    
    def generate_template_interview_prep(self, resume_text: str, job_description: str, interview_type: str) -> Dict[str, Any]:
        """
        Generate template-based interview preparation (fallback)
        """
        prep_data = {
            'questions': self.get_fallback_questions(interview_type),
            'tips': self.get_fallback_tips(interview_type),
            'examples': []
        }
        return prep_data
    
    def get_fallback_questions(self, interview_type: str) -> List[str]:
        """Get fallback questions by interview type"""
        if interview_type == 'behavioral':
            return [
                "Tell me about a time when you had to overcome a significant challenge.",
                "Describe a situation where you had to work with a difficult team member.",
                "Give me an example of when you had to learn something new quickly.",
                "Tell me about a time when you had to make a decision with limited information.",
                "Describe a project you're particularly proud of and why."
            ]
        elif interview_type == 'technical':
            return [
                "Explain the difference between abstract classes and interfaces.",
                "How would you optimize a slow-running database query?",
                "Describe your approach to debugging a production issue.",
                "What are the key principles of good software design?",
                "How do you ensure code quality in your projects?"
            ]
        elif interview_type == 'case':
            return [
                "How would you approach increasing user engagement for our product?",
                "Estimate the market size for electric vehicles in the next 5 years.",
                "Design a solution to reduce customer churn by 20%.",
                "How would you prioritize features for a new mobile app?",
                "Analyze the pros and cons of expanding into a new market."
            ]
        return []
    
    def get_fallback_tips(self, interview_type: str) -> List[str]:
        """Get fallback tips by interview type"""
        if interview_type == 'behavioral':
            return [
                "Use the STAR method (Situation, Task, Action, Result) to structure your answers.",
                "Prepare specific examples from your experience for each common behavioral question.",
                "Focus on your role and contributions in team scenarios.",
                "Quantify your achievements whenever possible.",
                "Practice your stories out loud to ensure they flow naturally."
            ]
        elif interview_type == 'technical':
            return [
                "Review fundamental concepts in your primary programming languages.",
                "Practice coding problems on platforms like LeetCode or HackerRank.",
                "Be prepared to explain your thought process while solving problems.",
                "Review system design concepts for senior-level positions.",
                "Prepare questions about the company's technology stack."
            ]
        elif interview_type == 'case':
            return [
                "Structure your approach clearly before diving into the analysis.",
                "Ask clarifying questions to understand the scope and constraints.",
                "Think out loud and explain your reasoning throughout the process.",
                "Consider multiple solutions and weigh their trade-offs.",
                "Practice with case study frameworks like MECE or Porter's Five Forces."
            ]
        return []
    
    def save_analysis(self, session_id: str, resume_text: str, job_description: str, industry: str, analysis: Dict[str, Any]):
        """Save analysis results to database"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute('''
                    INSERT INTO analyses (session_id, resume_text, job_description, industry, overall_score, analysis_data)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (session_id, resume_text, job_description, industry, analysis['overall_score'], json.dumps(analysis)))
                conn.commit()
        except Exception as e:
            logger.error(f"Failed to save analysis: {e}")
    
    def track_event(self, event_type: str, event_data: Dict[str, Any]):
        """Track analytics events"""
        try:
            session_id = session.get('session_id', 'anonymous')
            with sqlite3.connect(self.db_path) as conn:
                conn.execute('''
                    INSERT INTO analytics (event_type, event_data, session_id)
                    VALUES (?, ?, ?)
                ''', (event_type, json.dumps(event_data), session_id))
                conn.commit()
        except Exception as e:
            logger.error(f"Failed to track event: {e}")
    
    def generate_session_id(self) -> str:
        """Generate a unique session ID"""
        return hashlib.md5(f"{time.time()}{os.urandom(16)}".encode()).hexdigest()
    
    def run(self, host='127.0.0.1', port=None, debug=True):
        """Run the Flask application"""
        if port is None:
            port = int(os.getenv('PORT', 5000))
        logger.info(f"🌟 Starting AI Career Success Platform on http://{host}:{port}")
        self.app.run(host=host, port=port, debug=debug)

# Create application instance
career_platform = CareerSuccessPlatform()

# For deployment
app = career_platform.app

if __name__ == '__main__':
    # Run the application
    career_platform.run()
